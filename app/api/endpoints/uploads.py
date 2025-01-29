import base64
from io import BytesIO
import logging
import zipfile
from docx import Document
from fastapi import APIRouter, HTTPException
import os
from fastapi.responses import FileResponse
import openpyxl
import requests
from app.services.ects_service import get_ects_for_template
from app.services.prisma_service import get_template_from_prisma
from app.services.excel_service import match_template_and_get_word, process_excel_with_template
from app.services.word_service import generate_bulletins_from_excel
from app.services.ypareo_service import YpareoService
from prisma import Prisma
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches


router = APIRouter()

def calculate_weighted_average(notes):
    """
    Calcule la moyenne pondérée des notes avec leurs coefficients.
    Format des notes : "10 (0,25)" ou "10(0.25)" ou "10"
    """
    try:
        if not notes:
            return 0

        total_weighted_sum = 0
        total_coefficients = 0

        for note_str in notes:
            if note_str is None or str(note_str).strip() == "":
                continue

            note_str = str(note_str).strip()
            
            # Cas d'une note avec coefficient
            if "(" in note_str:
                try:
                    # Extraire la note et le coefficient
                    parts = note_str.split("(")
                    note = float(parts[0].strip())
                    coef = float(parts[1].replace(")", "").replace(",", ".").strip())
                    
                    total_weighted_sum += note * coef
                    total_coefficients += coef
                except (ValueError, IndexError):
                    continue
            
            # Cas d'une note simple (coefficient 1)
            else:
                try:
                    note = float(note_str)
                    total_weighted_sum += note
                    total_coefficients += 1
                except ValueError:
                    continue

        if total_coefficients == 0:
            return 0

        # Arrondir au centième
        return round(total_weighted_sum / total_coefficients, 2)

    except Exception as e:
        logging.error(f"Erreur lors du calcul de la moyenne pondérée: {str(e)}")
        return 0
    

def calculate_single_note_average(note_str):
    """
    Calcule la moyenne pour une note avec plusieurs coefficients.
    Format: "10 (0,25) - 15 (0,25) - 10,5 (0,5)" ou "17 - 16 - 17"
    """
    try:
        if not note_str or str(note_str).strip() == "":
            return ""

        note_str = str(note_str).strip()
        
        # Si la note contient des coefficients
        if "(" in note_str:
            try:
                # Séparer les différentes notes
                notes_parts = note_str.split("-")
                total_weighted_sum = 0
                total_coefficients = 0
                
                for part in notes_parts:
                    part = part.strip()
                    # Ignorer les parties contenant "Absent au devoir"
                    if "Absent au devoir" in part:
                        continue
                        
                    if "(" in part and ")" in part:
                        # Extraire la note et le coefficient
                        note_part = part.split("(")
                        # Remplacer la virgule par un point dans la note
                        note = float(note_part[0].strip().replace(",", "."))
                        # Remplacer la virgule par un point dans le coefficient
                        coeff = float(note_part[1].replace(",", ".").replace(")", "").strip())
                        
                        total_weighted_sum += note * coeff
                        total_coefficients += coeff
                
                if total_coefficients > 0:
                    # Arrondir à 2 décimales
                    return f"{(total_weighted_sum / total_coefficients):.2f}"
                return ""
                    
            except (ValueError, IndexError) as e:
                logging.error(f"Erreur lors du calcul de la moyenne pondérée: {str(e)}")
                return ""
        
        # Cas d'une note simple ou multiple sans coefficient
        else:
            try:
                # Séparer les notes s'il y en a plusieurs
                notes = [float(n.strip().replace(",", ".")) for n in note_str.split("-") if n.strip() and "Absent au devoir" not in n]
                if notes:
                    # Arrondir à 2 décimales
                    return f"{(sum(notes) / len(notes)):.2f}"
                return ""
            except ValueError as e:
                logging.error(f"Erreur lors du calcul de la note simple: {str(e)}")
                return ""

    except Exception as e:
        logging.error(f"Erreur lors du calcul de la note: {str(e)}")
        return ""
    

def calculate_ects_weighted_average(notes, ects_values):
    """
    Calcule la moyenne pondérée par les ECTS.
    Ne prend en compte que les matières avec des ECTS > 0.
    """
    try:
        total_weighted_sum = 0
        total_ects = 0

        # Parcourir les notes et leurs ECTS correspondants
        for note_str, ects_str in zip(notes, ects_values):
            if not note_str:
                continue

            try:
                # Convertir la note en float (remplacer la virgule par un point)
                note = float(str(note_str).replace(",", "."))
                # Convertir l'ECTS en entier
                ects = int(ects_str)
                
                # Toujours prendre en compte la note si elle existe, même avec ECTS = 0
                if note > 0:
                    # Si ECTS = 0, utiliser un coefficient de 1 pour la moyenne
                    coeff = max(ects, 1)
                    total_weighted_sum += note * coeff
                    total_ects += coeff

            except (ValueError, TypeError):
                continue

        if total_ects == 0:
            return ""

        # Arrondir à 2 décimales
        return f"{(total_weighted_sum / total_ects):.2f}"

    except Exception as e:
        logging.error(f"Erreur lors du calcul de la moyenne pondérée ECTS: {str(e)}")
        return ""


def get_etat(note_str: str, has_r_in_ue: bool = False) -> str:
    """
    Détermine l'état en fonction de la note.
    
    Args:
        note_str: La note sous forme de chaîne
        has_r_in_ue: Indique s'il y a déjà un R dans l'UE
    
    Returns:
        - Si note >= 10 : "VA"
        - Si note < 8 : "R"
        - Si 8 <= note < 10 : "C"
        - Si note est vide ou invalide : ""
    """
    if not note_str or str(note_str).strip() == "":
        return ""
        
    try:
        note = float(str(note_str).replace(",", "."))
        
        # Si la note est >= 10, c'est toujours "VA"
        if note >= 10:
            return "VA"
        # Si la note est < 8, c'est toujours "R"
        elif note < 8:
            return "R"
        # Si 8 <= note < 10, c'est "C"
        else:
            return "C"
    except (ValueError, TypeError):
        return ""



def get_etat_ue(etats: list, moyenne_ue: str = "") -> str:
    """
    Détermine l'état d'une UE en fonction des états des notes et de la moyenne de l'UE.
    
    Règles:
    - "VA" si moyenne_ue >= 10 ET pas de "R" dans les états
    - "NV" si:
        * au moins un état "R"
        * OU moyenne_ue < 10
    """
    try:
        moyenne = float(str(moyenne_ue).replace(",", ".")) if moyenne_ue else 0
    except (ValueError, TypeError):
        moyenne = 0

    # Vérifier s'il y a au moins un R
    has_r = "R" in etats
    
    # Si pas de R et moyenne >= 10 => VA
    if not has_r and moyenne >= 10:
        return "VA"
    
    # Sinon => NV
    return "NV"


    
def clean_temp_directory(temp_dir: str):
    for filename in os.listdir(temp_dir):
        file_path = os.path.join(temp_dir, filename)
        if os.path.isfile(file_path):
            os.remove(file_path)

def clean_except_specific_file(directory: str, keep_filename: str):
    """
    Supprime tous les fichiers dans un répertoire sauf le fichier spécifié.
    """
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        if filename != keep_filename and os.path.isfile(file_path):
            try:
                os.remove(file_path)
                logging.info(f"Fichier supprimé : {file_path}")
            except Exception as e:
                logging.error(f"Erreur lors de la suppression de {file_path}: {str(e)}")


# Calculer les totaux d'ECTS pour chaque UE en tenant compte de la règle des notes < 8
def calculate_ue_ects(notes, ects_values):
    total_ects = 0
    for note_str, ects_str in zip(notes, ects_values):
        try:
            note = float(note_str) if note_str else 0
            ects = int(ects_str)
            if note >= 8:  # On ne compte les ECTS que si la note est >= 8
                total_ects += ects
        except (ValueError, TypeError):
            continue
        return total_ects
    
def get_total_etat(etat_ue1: str, etat_ue2: str, etat_ue3: str, etat_ue4: str) -> str:
    """
    Détermine l'état total en fonction des états des UE.
    
    Règles:
    - "VA" si tous les états des UE sont "VA"
    - "NV" si au moins un état d'UE est "NV"
    """
    if all(etat == "VA" for etat in [etat_ue1, etat_ue2, etat_ue3, etat_ue4]):
        return "VA"
    return "NV"

@router.post("/process-excel")
async def process_excel(excel_url: str, word_url: str, user_id: str):
    try:
        output_dir = "./temp"

        logging.info(f"Début du traitement avec excel_url={excel_url}, word_url={word_url}")

        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        else:
            clean_temp_directory(output_dir)

        # Télécharger le fichier Excel source pour lire le nom du groupe
        logging.info(f"Téléchargement du fichier Excel depuis {excel_url}")
        excel_response = requests.get(excel_url)
        if excel_response.status_code != 200:
            raise HTTPException(status_code=400, detail="Impossible de télécharger le fichier Excel.")
        excel_file = BytesIO(excel_response.content)
        
        # Lire le nom du groupe depuis B2
        wb = openpyxl.load_workbook(excel_file)
        ws = wb.active
        group_name = ws["B2"].value
        
        if not group_name:
            raise HTTPException(status_code=400, detail="Nom du groupe non trouvé dans la cellule B2")
            
        # Déterminer le template à utiliser
        from app.core.template_mapping import TEMPLATE_MAPPING
        prisma_template = TEMPLATE_MAPPING.get(group_name)
        if not prisma_template:
            raise HTTPException(status_code=400, detail=f"Aucun template trouvé pour le groupe : {group_name}")

        logging.info(f"Template sélectionné pour le groupe {group_name}: {prisma_template}")

        # Récupérer le template Excel depuis Prisma
        template_excel_path = await get_template_from_prisma(prisma_template, output_dir)
        logging.info(f"Template récupéré et sauvegardé dans : {template_excel_path}")

        clean_except_specific_file(output_dir, keep_filename=os.path.basename(template_excel_path))

        logging.info("Début du traitement des données entre fichier source et template")
        result = await process_excel_with_template(excel_url, output_dir, prisma_template, user_id)
        logging.info(f"Traitement terminé. Fichier mis à jour disponible : {result['excel_path']}")

        return {"message": "Fichier traité avec succès", "excel_id": result['excel_id']}

    except Exception as e:
        logging.error(f"Erreur pendant le traitement : {str(e)}")
        raise HTTPException(status_code=400, detail=f"Erreur : {str(e)}")



@router.post("/get-word-template")
async def get_word_template_endpoint():
    try:
        db = Prisma()
        await db.connect()

        excel_path = os.path.join("./temp", "updated_excel.xlsx")
        if not os.path.exists(excel_path):
            raise ValueError("Fichier Excel mis à jour non trouvé dans ./temp")

        # Déterminer le template à utiliser en comparant avec les modèles
        template_info = await match_template_and_get_word(excel_path)
        template_name = template_info["template_name"]
        ects_template = template_info["ects_template"]
        
        logging.info(f"Utilisation du template {template_name} avec ECTS {ects_template}")

        updated_wb = openpyxl.load_workbook(excel_path)
        updated_ws = updated_wb.active

        date_du_jour = datetime.utcnow().strftime("%d/%m/%Y")

        # Mapping des colonnes selon le modèle
        if template_name == "modeleBG-ALT-S1-2024-2025.docx":
            ue_matieres = {
                "UE1_Title": str(updated_ws["C1"].value or ""),
                "matiere1": str(updated_ws["D1"].value or ""),
                "matiere2": str(updated_ws["E1"].value or ""),
                "matiere3": str(updated_ws["F1"].value or ""),
                "UE2_Title": str(updated_ws["G1"].value or ""),
                "matiere4": str(updated_ws["H1"].value or ""),
                "matiere5": str(updated_ws["I1"].value or ""),
                "matiere6": str(updated_ws["J1"].value or ""),
                "UE3_Title": str(updated_ws["K1"].value or ""),
                "matiere7": str(updated_ws["L1"].value or ""),
                "UE4_Title": str(updated_ws["M1"].value or ""),
                "matiere8": str(updated_ws["N1"].value or ""),
                "matiere9": str(updated_ws["O1"].value or ""),
                "matiere10": str(updated_ws["P1"].value or ""),
                "matiere11": str(updated_ws["Q1"].value or ""),
                "matiere12": str(updated_ws["R1"].value or ""),
                "matiere13": str(updated_ws["S1"].value or ""),
            }
        elif template_name == "modeleBG-ALT-S2-2024-2025.docx":  # modeleBGALT3.docx
            ue_matieres = {
                "UE1_Title": str(updated_ws["C1"].value or ""),
                "matiere1": str(updated_ws["D1"].value or ""),
                "matiere2": str(updated_ws["E1"].value or ""),
                "matiere3": str(updated_ws["F1"].value or ""),
                "matiere4": str(updated_ws["G1"].value or ""),
                "UE2_Title": str(updated_ws["H1"].value or ""),
                "matiere5": str(updated_ws["I1"].value or ""),
                "matiere6": str(updated_ws["J1"].value or ""),
                "matiere7": str(updated_ws["K1"].value or ""),
                "UE3_Title": str(updated_ws["L1"].value or ""),
                "matiere8": str(updated_ws["M1"].value or ""),
                "UE4_Title": str(updated_ws["N1"].value or ""),
                "matiere9": str(updated_ws["O1"].value or ""),
                "matiere10": str(updated_ws["P1"].value or ""),
                "matiere11": str(updated_ws["Q1"].value or ""),
                "matiere12": str(updated_ws["R1"].value or ""),
                "matiere13": str(updated_ws["S1"].value or ""),
            }
        elif template_name == "modeleBG-ALT-S3-2024-2025.docx":  # modeleBGALT3.docx
            ue_matieres = {
                "UE1_Title": str(updated_ws["C1"].value or ""),
                "matiere1": str(updated_ws["D1"].value or ""),
                "UE2_Title": str(updated_ws["E1"].value or ""),
                "matiere2": str(updated_ws["F1"].value or ""),
                "matiere3": str(updated_ws["G1"].value or ""),
                "matiere4": str(updated_ws["H1"].value or ""),
                "matiere5": str(updated_ws["I1"].value or ""),
                "UE3_Title": str(updated_ws["J1"].value or ""),
                "matiere6": str(updated_ws["K1"].value or ""),
                "matiere7": str(updated_ws["L1"].value or ""),
                "UE4_Title": str(updated_ws["M1"].value or ""),
                "matiere8": str(updated_ws["N1"].value or ""),
                "matiere9": str(updated_ws["O1"].value or ""),
                "matiere10": str(updated_ws["P1"].value or ""),
                "matiere11": str(updated_ws["Q1"].value or ""),
                "matiere12": str(updated_ws["R1"].value or ""),
            }
        elif template_name == "modeleBG-ALT-S4-2024-2025.docx":  # modeleBGALT3.docx
            ue_matieres = {
                "UE1_Title": str(updated_ws["C1"].value or ""),
                "matiere1": str(updated_ws["D1"].value or ""),
                "matiere2": str(updated_ws["E1"].value or ""),
                "matiere3": str(updated_ws["F1"].value or ""),
                "UE2_Title": str(updated_ws["G1"].value or ""),
                "matiere4": str(updated_ws["H1"].value or ""),
                "matiere5": str(updated_ws["I1"].value or ""),
                "matiere6": str(updated_ws["J1"].value or ""),
                "matiere7": str(updated_ws["K1"].value or ""),
                "matiere8": str(updated_ws["L1"].value or ""),
                "UE3_Title": str(updated_ws["M1"].value or ""),
                "matiere9": str(updated_ws["N1"].value or ""),
                "UE4_Title": str(updated_ws["O1"].value or ""),
                "matiere10": str(updated_ws["P1"].value or ""),
                "matiere11": str(updated_ws["Q1"].value or ""),
                "matiere12": str(updated_ws["R1"].value or ""),
                "matiere13": str(updated_ws["S1"].value or "")
            }
        elif template_name == "modeleBG-ALT-S5-2024-2025.docx":  # modeleBGALT3.docx
            ue_matieres = {
                "UE1_Title": str(updated_ws["C1"].value or ""),
                "matiere1": str(updated_ws["D1"].value or ""),
                "matiere2": str(updated_ws["E1"].value or ""),
                "UE2_Title": str(updated_ws["F1"].value or ""),
                "matiere3": str(updated_ws["G1"].value or ""),
                "matiere4": str(updated_ws["H1"].value or ""),
                "matiere5": str(updated_ws["I1"].value or ""),
                "matiere6": str(updated_ws["J1"].value or ""),
                "UE3_Title": str(updated_ws["K1"].value or ""),
                "matiere7": str(updated_ws["L1"].value or ""),
                "matiere8": str(updated_ws["M1"].value or ""),
                "UE4_Title": str(updated_ws["N1"].value or ""),
                "matiere9": str(updated_ws["O1"].value or ""),
                "matiere10": str(updated_ws["P1"].value or ""),
                "matiere11": str(updated_ws["Q1"].value or ""),
                "matiere12": str(updated_ws["R1"].value or ""),
                "matiere13": str(updated_ws["S1"].value or ""),
                "matiere14": str(updated_ws["T1"].value or "")
            }
        elif template_name == "modeleBG-ALT-S6-2024-2025.docx":  # modeleBGALT3.docx
            ue_matieres = {
                "UE1_Title": str(updated_ws["C1"].value or ""),
                "matiere1": str(updated_ws["D1"].value or ""),
                "matiere2": str(updated_ws["E1"].value or ""),
                "UE2_Title": str(updated_ws["F1"].value or ""),
                "matiere3": str(updated_ws["G1"].value or ""),
                "matiere4": str(updated_ws["H1"].value or ""),
                "UE3_Title": str(updated_ws["I1"].value or ""),
                "matiere5": str(updated_ws["J1"].value or ""),
                "matiere6": str(updated_ws["K1"].value or ""),
                "matiere7": str(updated_ws["L1"].value or ""),
                "UE4_Title": str(updated_ws["L1"].value or ""),
                "matiere8": str(updated_ws["M1"].value or ""),
                "matiere9": str(updated_ws["N1"].value or ""),
                "matiere10": str(updated_ws["O1"].value or ""),
                "matiere11": str(updated_ws["P1"].value or ""),
                "matiere12": str(updated_ws["Q1"].value or ""),
                "matiere13": str(updated_ws["R1"].value or "")
            }
        elif template_name == "modeleBG-TP-S1-2024-2025.docx":
            ue_matieres = {
                "UE1_Title": str(updated_ws["C1"].value or ""),
                "matiere1": str(updated_ws["D1"].value or ""),
                "matiere2": str(updated_ws["E1"].value or ""),
                "matiere3": str(updated_ws["F1"].value or ""),
                "matiere4 ": str(updated_ws["G1"].value or ""),
                "matiere5": str(updated_ws["H1"].value or ""),
                "matiere6": str(updated_ws["I1"].value or ""),
                "matiere7": str(updated_ws["J1"].value or ""),
                "UE2_Title": str(updated_ws["K1"].value or ""),
                "matiere8": str(updated_ws["L1"].value or ""),
                "matiere9": str(updated_ws["M1"].value or ""),
                "matiere10": str(updated_ws["N1"].value or ""),
                "matiere11": str(updated_ws["O1"].value or ""),
                "matiere12": str(updated_ws["P1"].value or ""),
                "matiere13": str(updated_ws["Q1"].value or ""),
                "UE3_Title": str(updated_ws["R1"].value or ""),
                "matiere14": str(updated_ws["S1"].value or ""),
                "matiere15": str(updated_ws["T1"].value or ""),
                "UE4_Title": str(updated_ws["U1"].value or ""),
                "matiere16": str(updated_ws["V1"].value or ""),
                "matiere17": str(updated_ws["W1"].value or ""),
                "matiere18": str(updated_ws["Z1"].value or ""),
                "matiere19": str(updated_ws["Y1"].value or ""),
                "matiere20": str(updated_ws["Z1"].value or ""),
            }
        elif template_name == "modeleBG-TP-S2-2024-2025.docx":  # modeleBGALT3.docx
            ue_matieres = {
                "UE1_Title": str(updated_ws["C1"].value or ""),
                "matiere1": str(updated_ws["D1"].value or ""),
                "matiere2": str(updated_ws["E1"].value or "")
            }
        elif template_name == "modeleBG-TP-S3-2024-2025.docx":  # modeleBGALT3.docx
            ue_matieres = {
                "UE1_Title": str(updated_ws["C1"].value or ""),
                "matiere1": str(updated_ws["D1"].value or ""),
                "UE2_Title": str(updated_ws["E1"].value or ""),
                "matiere2": str(updated_ws["F1"].value or ""),
                "matiere3": str(updated_ws["G1"].value or ""),
                "matiere4 ": str(updated_ws["H1"].value or ""),
                "matiere5": str(updated_ws["I1"].value or ""),
                "matiere6": str(updated_ws["J1"].value or ""),
                "matiere7": str(updated_ws["K1"].value or ""),
                "matiere8": str(updated_ws["L1"].value or ""),
                "matiere9": str(updated_ws["M1"].value or ""),
                "UE3_Title": str(updated_ws["N1"].value or ""),
                "matiere10": str(updated_ws["O1"].value or ""),
                "matiere11": str(updated_ws["P1"].value or ""),
                "UE4_Title": str(updated_ws["Q1"].value or ""),
                "matiere12": str(updated_ws["R1"].value or ""),
                "matiere13": str(updated_ws["S1"].value or ""),
                "matiere14": str(updated_ws["T1"].value or ""),
                "matiere15": str(updated_ws["U1"].value or ""),
            }
        elif template_name == "modeleBG-TP-S4-2024-2025.docx":  # modeleBGALT3.docx
            ue_matieres = {
                "UE1_Title": str(updated_ws["C1"].value or ""),
                "matiere1": str(updated_ws["D1"].value or "")
            }
        elif template_name == "modeleBG-TP-S5-2024-2025.docx":  # modeleBGALT3.docx
            ue_matieres = {
                "UE1_Title": str(updated_ws["C1"].value or ""),
                "matiere1": str(updated_ws["D1"].value or ""),
                "matiere2": str(updated_ws["E1"].value or ""),
                "matiere3": str(updated_ws["F1"].value or ""),
                "matiere4 ": str(updated_ws["G1"].value or ""),
                "matiere5": str(updated_ws["H1"].value or ""),
                "UE2_Title": str(updated_ws["I1"].value or ""),
                "matiere6": str(updated_ws["J1"].value or ""),
                "matiere7": str(updated_ws["K1"].value or ""),
                "matiere8": str(updated_ws["L1"].value or ""),
                "matiere9": str(updated_ws["M1"].value or ""),
                "matiere10": str(updated_ws["N1"].value or ""),
                "UE3_Title": str(updated_ws["O1"].value or ""),
                "matiere11": str(updated_ws["P1"].value or ""),
                "matiere12": str(updated_ws["Q1"].value or ""),
                "matiere13": str(updated_ws["R1"].value or ""),
                "matiere14": str(updated_ws["S1"].value or ""),
                "matiere15": str(updated_ws["T1"].value or ""),
                "UE4_Title": str(updated_ws["U1"].value or ""),
                "matiere16": str(updated_ws["V1"].value or ""),
                "matiere17": str(updated_ws["W1"].value or ""),
                "matiere18": str(updated_ws["X1"].value or ""),
                "matiere19": str(updated_ws["Y1"].value or ""),
                "matiere20": str(updated_ws["Z1"].value or "")
            }
        elif template_name == "modeleBG-TP-S6-2024-2025.docx":  # modeleBGALT3.docx
            ue_matieres = {
                "UE1_Title": str(updated_ws["C1"].value or ""),
                "matiere1": str(updated_ws["D1"].value or ""),
                "matiere2": str(updated_ws["E1"].value or ""),
                "matiere3": str(updated_ws["F1"].value or "")
            }
        elif template_name == "modeleM1-S1.docx":  # modeleBGALT3.docx
            ue_matieres = {
                "UE1_Title": str(updated_ws["C1"].value or ""),
                "matiere1": str(updated_ws["D1"].value or ""),
                "matiere2": str(updated_ws["E1"].value or ""),
                "UE2_Title": str(updated_ws["F1"].value or ""),
                "matiere3": str(updated_ws["G1"].value or ""),
                "matiere4 ": str(updated_ws["H1"].value or ""),
                "UE3_Title": str(updated_ws["I1"].value or ""),
                "matiere5": str(updated_ws["J1"].value or ""),
                "matiere6": str(updated_ws["K1"].value or ""),
                "UE4_Title": str(updated_ws["L1"].value or ""),
                "matiere7": str(updated_ws["M1"].value or ""),
                "matiere8": str(updated_ws["O1"].value or ""),
                "matiere9": str(updated_ws["P1"].value or ""),
                "matiere10": str(updated_ws["Q1"].value or ""),
                "matiere11": str(updated_ws["R1"].value or ""),
                "UESPE_Title": str(updated_ws["S1"].value or ""),
                "matiere12": str(updated_ws["T1"].value or ""),
                "matiere13": str(updated_ws["U1"].value or ""),
                "matiere14": str(updated_ws["V1"].value or ""),
            }
        elif template_name == "modeleM2-S3.docx":  # modeleBGALT3.docx
            ue_matieres = {
                "UE1_Title": str(updated_ws["C1"].value or ""),
                "matiere1": str(updated_ws["D1"].value or ""),
                "matiere2": str(updated_ws["E1"].value or ""),
                "UE2_Title": str(updated_ws["F1"].value or ""),
                "matiere3": str(updated_ws["G1"].value or ""),
                "UE3_Title": str(updated_ws["H1"].value or ""),
                "matiere4 ": str(updated_ws["I1"].value or ""),
                "UE4_Title": str(updated_ws["J1"].value or ""),
                "matiere5": str(updated_ws["K1"].value or ""),
                "matiere6": str(updated_ws["L1"].value or ""),
                "matiere7": str(updated_ws["M1"].value or ""),
                "matiere8": str(updated_ws["N1"].value or ""),
                "matiere9": str(updated_ws["O1"].value or ""),
                "UESPE_Title": str(updated_ws["P1"].value or ""),
                "matiere10": str(updated_ws["Q1"].value or ""),
                "matiere11": str(updated_ws["R1"].value or ""),
                "matiere12": str(updated_ws["S1"].value or ""),
                "matiere13": str(updated_ws["T1"].value or ""),
            }

        # Récupérer le template Word
        word_template = await db.generatedfile.find_first(
            where={
                "filename": template_name,
                "isTemplate": True
            }
        )
        
        if not word_template:
            raise ValueError(f"Template Word {template_name} non trouvé dans Prisma")

        # Récupérer les ECTS selon le template
        ects_data = await get_ects_for_template(ects_template)
        logging.info(f"ECTS data for {ects_template}: {ects_data}")
        
        if template_name == "modeleBG-ALT-S1-2024-2025.docx":
            required_ects = [f"ECTS{i}" for i in range(1, 14)]  # ECTS1 à ECTS13
            missing_ects = [ects for ects in required_ects if ects not in ects_data]
            if missing_ects:
                raise ValueError(f"Missing ECTS values for {template_name}: {missing_ects}")

        elif template_name == "modeleBG-ALT-S2-2024-2025.docx":
            required_ects = [f"ECTS{i}" for i in range(1, 14)]  # ECTS1 à ECTS13
            missing_ects = [ects for ects in required_ects if ects not in ects_data]
            if missing_ects:
                raise ValueError(f"Missing ECTS values for {template_name}: {missing_ects}")

        elif template_name == "modeleBG-ALT-S3-2024-2025.docx":
            required_ects = [f"ECTS{i}" for i in range(1, 13)]  # ECTS1 à ECTS12
            missing_ects = [ects for ects in required_ects if ects not in ects_data]
            if missing_ects:
                raise ValueError(f"Missing ECTS values for {template_name}: {missing_ects}")

        elif template_name == "modeleBG-ALT-S4-2024-2025.docx":
            required_ects = [f"ECTS{i}" for i in range(1, 14)]  # ECTS1 à ECTS13
            missing_ects = [ects for ects in required_ects if ects not in ects_data]
            if missing_ects:
                raise ValueError(f"Missing ECTS values for {template_name}: {missing_ects}")

        elif template_name == "modeleBG-ALT-S5-2024-2025.docx":
            required_ects = [f"ECTS{i}" for i in range(1, 15)]  # ECTS1 à ECTS14
            missing_ects = [ects for ects in required_ects if ects not in ects_data]
            if missing_ects:
                raise ValueError(f"Missing ECTS values for {template_name}: {missing_ects}")

        elif template_name == "modeleBG-ALT-S6-2024-2025.docx":
            required_ects = [f"ECTS{i}" for i in range(1, 14)]  # ECTS1 à ECTS13
            missing_ects = [ects for ects in required_ects if ects not in ects_data]
            if missing_ects:
                raise ValueError(f"Missing ECTS values for {template_name}: {missing_ects}")

        elif template_name == "modeleBG-TP-S1-2024-2025.docx":
            required_ects = [f"ECTS{i}" for i in range(1, 21)]  # ECTS1 à ECTS20
            missing_ects = [ects for ects in required_ects if ects not in ects_data]
            if missing_ects:
                raise ValueError(f"Missing ECTS values for {template_name}: {missing_ects}")

        elif template_name == "modeleBG-TP-S2-2024-2025.docx":
            required_ects = [f"ECTS{i}" for i in range(1, 3)]  # ECTS1 à ECTS2
            missing_ects = [ects for ects in required_ects if ects not in ects_data]
            if missing_ects:
                raise ValueError(f"Missing ECTS values for {template_name}: {missing_ects}")

        elif template_name == "modeleBG-TP-S3-2024-2025.docx":
            required_ects = [f"ECTS{i}" for i in range(1, 16)]  # ECTS1 à ECTS15
            missing_ects = [ects for ects in required_ects if ects not in ects_data]
            if missing_ects:
                raise ValueError(f"Missing ECTS values for {template_name}: {missing_ects}")

        elif template_name == "modeleBG-TP-S4-2024-2025.docx":
            required_ects = [f"ECTS{i}" for i in range(1, 2)]  # ECTS1
            missing_ects = [ects for ects in required_ects if ects not in ects_data]
            if missing_ects:
                raise ValueError(f"Missing ECTS values for {template_name}: {missing_ects}")

        elif template_name == "modeleBG-TP-S5-2024-2025.docx":
            required_ects = [f"ECTS{i}" for i in range(1, 21)]  # ECTS1 à ECTS20
            missing_ects = [ects for ects in required_ects if ects not in ects_data]
            if missing_ects:
                raise ValueError(f"Missing ECTS values for {template_name}: {missing_ects}")

        elif template_name == "modeleBG-TP-S6-2024-2025.docx":
            required_ects = [f"ECTS{i}" for i in range(1, 4)]  # ECTS1 à ECTS3
            missing_ects = [ects for ects in required_ects if ects not in ects_data]
            if missing_ects:
                raise ValueError(f"Missing ECTS values for {template_name}: {missing_ects}")

        elif template_name == "modeleM1-S1.docx":
            required_ects = [f"ECTS{i}" for i in range(1, 15)]  # ECTS1 à ECTS14
            missing_ects = [ects for ects in required_ects if ects not in ects_data]
            if missing_ects:
                raise ValueError(f"Missing ECTS values for {template_name}: {missing_ects}")

        elif template_name == "modeleM2-S3.docx":
            required_ects = [f"ECTS{i}" for i in range(1, 14)]  # ECTS1 à ECTS13
            missing_ects = [ects for ects in required_ects if ects not in ects_data]
            if missing_ects:
                raise ValueError(f"Missing ECTS values for {template_name}: {missing_ects}")

            required_ects = [f"ECTS{i}" for i in range(1, 14)]  # ECTS1 à ECTS13
            missing_ects = [ects for ects in required_ects if ects not in ects_data]
            if missing_ects:
                raise ValueError(f"Missing ECTS values for {template_name}: {missing_ects}")

        bulletins_dir = os.path.join("./temp", "bulletins")
        if not os.path.exists(bulletins_dir):
            os.makedirs(bulletins_dir)

        for row in range(3, updated_ws.max_row + 1):
            if not updated_ws[f"B{row}"].value:
                continue

            # Récupérer les notes selon le modèle
            if template_name == "modeleBG-ALT-S1-2024-2025.docx":
                ue1_notes = [
                    calculate_single_note_average(updated_ws[f"D{row}"].value),
                    calculate_single_note_average(updated_ws[f"E{row}"].value),
                    calculate_single_note_average(updated_ws[f"F{row}"].value)
                ]
                ue2_notes = [
                    calculate_single_note_average(updated_ws[f"H{row}"].value),
                    calculate_single_note_average(updated_ws[f"I{row}"].value),
                    calculate_single_note_average(updated_ws[f"J{row}"].value)
                ]
                ue3_notes = [
                    calculate_single_note_average(updated_ws[f"L{row}"].value)
                ]
                ue4_notes = [
                    calculate_single_note_average(updated_ws[f"N{row}"].value),
                    calculate_single_note_average(updated_ws[f"O{row}"].value),
                    calculate_single_note_average(updated_ws[f"P{row}"].value),
                    calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    calculate_single_note_average(updated_ws[f"R{row}"].value),
                    calculate_single_note_average(updated_ws[f"S{row}"].value)
                ]

            elif template_name == "modeleBG-ALT-S2-2024-2025.docx":
                ue1_notes = [
                    calculate_single_note_average(updated_ws[f"D{row}"].value),
                    calculate_single_note_average(updated_ws[f"E{row}"].value),
                    calculate_single_note_average(updated_ws[f"F{row}"].value),
                    calculate_single_note_average(updated_ws[f"G{row}"].value)
                ]
                ue2_notes = [
                    calculate_single_note_average(updated_ws[f"I{row}"].value),
                    calculate_single_note_average(updated_ws[f"J{row}"].value),
                    calculate_single_note_average(updated_ws[f"K{row}"].value)
                ]
                ue3_notes = [
                    calculate_single_note_average(updated_ws[f"M{row}"].value)
                ]
                ue4_notes = [
                    calculate_single_note_average(updated_ws[f"O{row}"].value),
                    calculate_single_note_average(updated_ws[f"P{row}"].value),
                    calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    calculate_single_note_average(updated_ws[f"R{row}"].value),
                    calculate_single_note_average(updated_ws[f"S{row}"].value)
                ]

            elif template_name == "modeleBG-ALT-S3-2024-2025.docx":
                ue1_notes = [
                    calculate_single_note_average(updated_ws[f"D{row}"].value)
                ]
                ue2_notes = [
                    calculate_single_note_average(updated_ws[f"F{row}"].value),
                    calculate_single_note_average(updated_ws[f"G{row}"].value),
                    calculate_single_note_average(updated_ws[f"H{row}"].value),
                    calculate_single_note_average(updated_ws[f"I{row}"].value)
                ]
                ue3_notes = [
                    calculate_single_note_average(updated_ws[f"K{row}"].value),
                    calculate_single_note_average(updated_ws[f"L{row}"].value)
                ]
                ue4_notes = [
                    calculate_single_note_average(updated_ws[f"N{row}"].value),
                    calculate_single_note_average(updated_ws[f"O{row}"].value),
                    calculate_single_note_average(updated_ws[f"P{row}"].value),
                    calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    calculate_single_note_average(updated_ws[f"R{row}"].value)
                ]

            elif template_name == "modeleBG-ALT-S4-2024-2025.docx":
                ue1_notes = [
                    calculate_single_note_average(updated_ws[f"D{row}"].value),
                    calculate_single_note_average(updated_ws[f"E{row}"].value),
                    calculate_single_note_average(updated_ws[f"F{row}"].value)
                ]
                ue2_notes = [
                    calculate_single_note_average(updated_ws[f"H{row}"].value),
                    calculate_single_note_average(updated_ws[f"I{row}"].value),
                    calculate_single_note_average(updated_ws[f"J{row}"].value),
                    calculate_single_note_average(updated_ws[f"K{row}"].value),
                    calculate_single_note_average(updated_ws[f"L{row}"].value)
                ]
                ue3_notes = [
                    calculate_single_note_average(updated_ws[f"N{row}"].value)
                ]
                ue4_notes = [
                    calculate_single_note_average(updated_ws[f"P{row}"].value),
                    calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    calculate_single_note_average(updated_ws[f"R{row}"].value),
                    calculate_single_note_average(updated_ws[f"S{row}"].value)
                ]

            elif template_name == "modeleBG-ALT-S5-2024-2025.docx":
                ue1_notes = [
                    calculate_single_note_average(updated_ws[f"D{row}"].value),
                    calculate_single_note_average(updated_ws[f"E{row}"].value)
                ]
                ue2_notes = [
                    calculate_single_note_average(updated_ws[f"G{row}"].value),
                    calculate_single_note_average(updated_ws[f"H{row}"].value),
                    calculate_single_note_average(updated_ws[f"I{row}"].value),
                    calculate_single_note_average(updated_ws[f"J{row}"].value)
                ]
                ue3_notes = [
                    calculate_single_note_average(updated_ws[f"L{row}"].value),
                    calculate_single_note_average(updated_ws[f"M{row}"].value)
                ]
                ue4_notes = [
                    calculate_single_note_average(updated_ws[f"O{row}"].value),
                    calculate_single_note_average(updated_ws[f"P{row}"].value),
                    calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    calculate_single_note_average(updated_ws[f"R{row}"].value),
                    calculate_single_note_average(updated_ws[f"S{row}"].value),
                    calculate_single_note_average(updated_ws[f"T{row}"].value)
                ]

            elif template_name == "modeleBG-ALT-S6-2024-2025.docx":
                ue1_notes = [
                    calculate_single_note_average(updated_ws[f"D{row}"].value),
                    calculate_single_note_average(updated_ws[f"E{row}"].value)
                ]
                ue2_notes = [
                    calculate_single_note_average(updated_ws[f"G{row}"].value),
                    calculate_single_note_average(updated_ws[f"H{row}"].value)
                ]
                ue3_notes = [
                    calculate_single_note_average(updated_ws[f"J{row}"].value),
                    calculate_single_note_average(updated_ws[f"K{row}"].value),
                    calculate_single_note_average(updated_ws[f"L{row}"].value)
                ]
                ue4_notes = [
                    calculate_single_note_average(updated_ws[f"N{row}"].value),
                    calculate_single_note_average(updated_ws[f"O{row}"].value),
                    calculate_single_note_average(updated_ws[f"P{row}"].value),
                    calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    calculate_single_note_average(updated_ws[f"R{row}"].value)
                ]

            elif template_name == "modeleBG-TP-S1-2024-2025.docx":
                ue1_notes = [
                    calculate_single_note_average(updated_ws[f"D{row}"].value),
                    calculate_single_note_average(updated_ws[f"E{row}"].value),
                    calculate_single_note_average(updated_ws[f"F{row}"].value),
                    calculate_single_note_average(updated_ws[f"G{row}"].value),
                    calculate_single_note_average(updated_ws[f"H{row}"].value),
                    calculate_single_note_average(updated_ws[f"I{row}"].value),
                    calculate_single_note_average(updated_ws[f"J{row}"].value)
                ]
                ue2_notes = [
                    calculate_single_note_average(updated_ws[f"L{row}"].value),
                    calculate_single_note_average(updated_ws[f"M{row}"].value),
                    calculate_single_note_average(updated_ws[f"N{row}"].value),
                    calculate_single_note_average(updated_ws[f"O{row}"].value),
                    calculate_single_note_average(updated_ws[f"P{row}"].value),
                    calculate_single_note_average(updated_ws[f"Q{row}"].value)
                ]
                ue3_notes = [
                    calculate_single_note_average(updated_ws[f"S{row}"].value),
                    calculate_single_note_average(updated_ws[f"T{row}"].value)
                ]
                ue4_notes = [
                    calculate_single_note_average(updated_ws[f"V{row}"].value),
                    calculate_single_note_average(updated_ws[f"W{row}"].value),
                    calculate_single_note_average(updated_ws[f"X{row}"].value),
                    calculate_single_note_average(updated_ws[f"Y{row}"].value),
                    calculate_single_note_average(updated_ws[f"Z{row}"].value)
                ]

            elif template_name == "modeleBG-TP-S2-2024-2025.docx":
                ue1_notes = [
                    calculate_single_note_average(updated_ws[f"D{row}"].value),
                    calculate_single_note_average(updated_ws[f"E{row}"].value)
                ]
                ue2_notes = []
                ue3_notes = []
                ue4_notes = []

            elif template_name == "modeleBG-TP-S3-2024-2025.docx":
                ue1_notes = [
                    calculate_single_note_average(updated_ws[f"D{row}"].value)
                ]
                ue2_notes = [
                    calculate_single_note_average(updated_ws[f"F{row}"].value),
                    calculate_single_note_average(updated_ws[f"G{row}"].value),
                    calculate_single_note_average(updated_ws[f"H{row}"].value),
                    calculate_single_note_average(updated_ws[f"I{row}"].value),
                    calculate_single_note_average(updated_ws[f"J{row}"].value),
                    calculate_single_note_average(updated_ws[f"K{row}"].value),
                    calculate_single_note_average(updated_ws[f"L{row}"].value),
                    calculate_single_note_average(updated_ws[f"M{row}"].value)
                ]
                ue3_notes = [
                    calculate_single_note_average(updated_ws[f"O{row}"].value),
                    calculate_single_note_average(updated_ws[f"P{row}"].value)
                ]
                ue4_notes = [
                    calculate_single_note_average(updated_ws[f"R{row}"].value),
                    calculate_single_note_average(updated_ws[f"S{row}"].value),
                    calculate_single_note_average(updated_ws[f"T{row}"].value),
                    calculate_single_note_average(updated_ws[f"U{row}"].value)
                ]

            elif template_name == "modeleBG-TP-S4-2024-2025.docx":
                ue1_notes = [
                    calculate_single_note_average(updated_ws[f"D{row}"].value)
                ]
                ue2_notes = []
                ue3_notes = []
                ue4_notes = []

            elif template_name == "modeleBG-TP-S5-2024-2025.docx":
                ue1_notes = [
                    calculate_single_note_average(updated_ws[f"D{row}"].value),
                    calculate_single_note_average(updated_ws[f"E{row}"].value),
                    calculate_single_note_average(updated_ws[f"F{row}"].value),
                    calculate_single_note_average(updated_ws[f"G{row}"].value),
                    calculate_single_note_average(updated_ws[f"H{row}"].value)
                ]
                ue2_notes = [
                    calculate_single_note_average(updated_ws[f"J{row}"].value),
                    calculate_single_note_average(updated_ws[f"K{row}"].value),
                    calculate_single_note_average(updated_ws[f"L{row}"].value),
                    calculate_single_note_average(updated_ws[f"M{row}"].value),
                    calculate_single_note_average(updated_ws[f"N{row}"].value)
                ]
                ue3_notes = [
                    calculate_single_note_average(updated_ws[f"P{row}"].value),
                    calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    calculate_single_note_average(updated_ws[f"R{row}"].value),
                    calculate_single_note_average(updated_ws[f"S{row}"].value),
                    calculate_single_note_average(updated_ws[f"T{row}"].value)
                ]
                ue4_notes = [
                    calculate_single_note_average(updated_ws[f"V{row}"].value),
                    calculate_single_note_average(updated_ws[f"W{row}"].value),
                    calculate_single_note_average(updated_ws[f"X{row}"].value),
                    calculate_single_note_average(updated_ws[f"Y{row}"].value),
                    calculate_single_note_average(updated_ws[f"Z{row}"].value)
                ]

            elif template_name == "modeleBG-TP-S6-2024-2025.docx":
                ue1_notes = [
                    calculate_single_note_average(updated_ws[f"D{row}"].value),
                    calculate_single_note_average(updated_ws[f"E{row}"].value),
                    calculate_single_note_average(updated_ws[f"F{row}"].value)
                ]
                ue2_notes = []
                ue3_notes = []
                ue4_notes = []

            elif template_name == "modeleM1-S1.docx":
                ue1_notes = [
                    calculate_single_note_average(updated_ws[f"D{row}"].value),
                    calculate_single_note_average(updated_ws[f"E{row}"].value)
                ]
                ue2_notes = [
                    calculate_single_note_average(updated_ws[f"G{row}"].value),
                    calculate_single_note_average(updated_ws[f"H{row}"].value)
                ]
                ue3_notes = [
                    calculate_single_note_average(updated_ws[f"J{row}"].value),
                    calculate_single_note_average(updated_ws[f"K{row}"].value)
                ]
                ue4_notes = [
                    calculate_single_note_average(updated_ws[f"M{row}"].value),
                    calculate_single_note_average(updated_ws[f"O{row}"].value),
                    calculate_single_note_average(updated_ws[f"P{row}"].value),
                    calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    calculate_single_note_average(updated_ws[f"R{row}"].value)
                ]
                uespe_notes = [
                    calculate_single_note_average(updated_ws[f"T{row}"].value),
                    calculate_single_note_average(updated_ws[f"U{row}"].value),
                    calculate_single_note_average(updated_ws[f"V{row}"].value)
                ]

            elif template_name == "modeleM2-S3.docx":
                ue1_notes = [
                    calculate_single_note_average(updated_ws[f"D{row}"].value),
                    calculate_single_note_average(updated_ws[f"E{row}"].value)
                ]
                ue2_notes = [
                    calculate_single_note_average(updated_ws[f"G{row}"].value)
                ]
                ue3_notes = [
                    calculate_single_note_average(updated_ws[f"I{row}"].value)
                ]
                ue4_notes = [
                    calculate_single_note_average(updated_ws[f"K{row}"].value),
                    calculate_single_note_average(updated_ws[f"L{row}"].value),
                    calculate_single_note_average(updated_ws[f"M{row}"].value),
                    calculate_single_note_average(updated_ws[f"N{row}"].value),
                    calculate_single_note_average(updated_ws[f"O{row}"].value)
                ]
                uespe_notes = [
                    calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    calculate_single_note_average(updated_ws[f"R{row}"].value),
                    calculate_single_note_average(updated_ws[f"S{row}"].value),
                    calculate_single_note_average(updated_ws[f"T{row}"].value)
                ]


            # Calculer les moyennes avec ECTS
            if template_name == "modeleBG-ALT-S1-2024-2025.docx":
                moyUE1 = calculate_ects_weighted_average(ue1_notes, [
                    ects_data["ECTS1"], ects_data["ECTS2"], ects_data["ECTS3"]
                ])
                moyUE2 = calculate_ects_weighted_average(ue2_notes, [
                    ects_data["ECTS4"], ects_data["ECTS5"], ects_data["ECTS6"]
                ])
                moyUE3 = calculate_ects_weighted_average(ue3_notes, [
                    ects_data["ECTS7"]
                ])
                moyUE4 = calculate_ects_weighted_average(ue4_notes, [
                    ects_data["ECTS8"], ects_data["ECTS9"], ects_data["ECTS10"],
                    ects_data["ECTS11"], ects_data["ECTS12"], ects_data["ECTS13"]
                ])
            
            # Calculer les totaux d'ECTS pour chaque UE
            if template_name == "modeleBG-ALT-S1-2024-2025.docx":
                ects_ue1 = sum(int(ects_data[f"ECTS{i}"]) for i in range(1, 4))
                ects_ue2 = sum(int(ects_data[f"ECTS{i}"]) for i in range(4, 7))
                ects_ue3 = int(ects_data["ECTS7"])
                ects_ue4 = sum(int(ects_data[f"ECTS{i}"]) for i in range(8, 14))

            elif template_name == "modeleBG-ALT-S2-2024-2025.docx":
                ects_ue1 = sum(int(ects_data[f"ECTS{i}"]) for i in range(1, 5))
                ects_ue2 = sum(int(ects_data[f"ECTS{i}"]) for i in range(5, 8))
                ects_ue3 = int(ects_data["ECTS8"])
                ects_ue4 = sum(int(ects_data[f"ECTS{i}"]) for i in range(9, 14))

            elif template_name == "modeleBG-ALT-S3-2024-2025.docx":
                ects_ue1 = int(ects_data["ECTS1"])
                ects_ue2 = sum(int(ects_data[f"ECTS{i}"]) for i in range(2, 6))
                ects_ue3 = sum(int(ects_data[f"ECTS{i}"]) for i in range(6, 8))
                ects_ue4 = sum(int(ects_data[f"ECTS{i}"]) for i in range(8, 13))

            elif template_name == "modeleBG-ALT-S4-2024-2025.docx":
                ects_ue1 = sum(int(ects_data[f"ECTS{i}"]) for i in range(1, 4))
                ects_ue2 = sum(int(ects_data[f"ECTS{i}"]) for i in range(4, 9))
                ects_ue3 = int(ects_data["ECTS9"])
                ects_ue4 = sum(int(ects_data[f"ECTS{i}"]) for i in range(10, 14))

            elif template_name == "modeleBG-ALT-S5-2024-2025.docx":
                ects_ue1 = sum(int(ects_data[f"ECTS{i}"]) for i in range(1, 3))
                ects_ue2 = sum(int(ects_data[f"ECTS{i}"]) for i in range(3, 7))
                ects_ue3 = sum(int(ects_data[f"ECTS{i}"]) for i in range(7, 9))
                ects_ue4 = sum(int(ects_data[f"ECTS{i}"]) for i in range(9, 15))

            elif template_name == "modeleBG-ALT-S6-2024-2025.docx":
                ects_ue1 = sum(int(ects_data[f"ECTS{i}"]) for i in range(1, 3))
                ects_ue2 = sum(int(ects_data[f"ECTS{i}"]) for i in range(3, 5))
                ects_ue3 = sum(int(ects_data[f"ECTS{i}"]) for i in range(5, 8))
                ects_ue4 = sum(int(ects_data[f"ECTS{i}"]) for i in range(8, 14))

            elif template_name == "modeleBG-TP-S1-2024-2025.docx":
                ects_ue1 = sum(int(ects_data[f"ECTS{i}"]) for i in range(1, 8))
                ects_ue2 = sum(int(ects_data[f"ECTS{i}"]) for i in range(8, 14))
                ects_ue3 = sum(int(ects_data[f"ECTS{i}"]) for i in range(14, 16))
                ects_ue4 = sum(int(ects_data[f"ECTS{i}"]) for i in range(16, 21))

            elif template_name == "modeleBG-TP-S2-2024-2025.docx":
                ects_ue1 = sum(int(ects_data[f"ECTS{i}"]) for i in range(1, 3))
                ects_ue2 = 0
                ects_ue3 = 0
                ects_ue4 = 0

            elif template_name == "modeleBG-TP-S3-2024-2025.docx":
                ects_ue1 = int(ects_data["ECTS1"])
                ects_ue2 = sum(int(ects_data[f"ECTS{i}"]) for i in range(2, 10))
                ects_ue3 = sum(int(ects_data[f"ECTS{i}"]) for i in range(10, 12))
                ects_ue4 = sum(int(ects_data[f"ECTS{i}"]) for i in range(12, 16))

            elif template_name == "modeleBG-TP-S4-2024-2025.docx":
                ects_ue1 = int(ects_data["ECTS1"])
                ects_ue2 = 0
                ects_ue3 = 0
                ects_ue4 = 0

            elif template_name == "modeleBG-TP-S5-2024-2025.docx":
                ects_ue1 = sum(int(ects_data[f"ECTS{i}"]) for i in range(1, 6))
                ects_ue2 = sum(int(ects_data[f"ECTS{i}"]) for i in range(6, 11))
                ects_ue3 = sum(int(ects_data[f"ECTS{i}"]) for i in range(11, 16))
                ects_ue4 = sum(int(ects_data[f"ECTS{i}"]) for i in range(16, 21))

            elif template_name == "modeleBG-TP-S6-2024-2025.docx":
                ects_ue1 = sum(int(ects_data[f"ECTS{i}"]) for i in range(1, 4))
                ects_ue2 = 0
                ects_ue3 = 0
                ects_ue4 = 0

            elif template_name == "modeleM1-S1.docx":
                ects_ue1 = sum(int(ects_data[f"ECTS{i}"]) for i in range(1, 3))
                ects_ue2 = sum(int(ects_data[f"ECTS{i}"]) for i in range(3, 5))
                ects_ue3 = sum(int(ects_data[f"ECTS{i}"]) for i in range(5, 7))
                ects_ue4 = sum(int(ects_data[f"ECTS{i}"]) for i in range(7, 12))
                ects_uespe = sum(int(ects_data[f"ECTS{i}"]) for i in range(12, 15))

            elif template_name == "modeleM2-S3.docx":
                ects_ue1 = sum(int(ects_data[f"ECTS{i}"]) for i in range(1, 3))
                ects_ue2 = int(ects_data["ECTS3"])
                ects_ue3 = int(ects_data["ECTS4"])
                ects_ue4 = sum(int(ects_data[f"ECTS{i}"]) for i in range(5, 10))
                ects_uespe = sum(int(ects_data[f"ECTS{i}"]) for i in range(10, 14))

            moyenne_ects = ects_ue1 + ects_ue2 + ects_ue3 + ects_ue4
            if template_name in ["modeleM1-S1.docx", "modeleM2-S3.docx"]:
                moyenne_ects += ects_uespe

            try:
                if moyenne_ects > 0:
                    if template_name in ["modeleM1-S1.docx", "modeleM2-S3.docx"]:
                        moyenne_ponderee = (
                            float(moyUE1 or 0) * ects_ue1 +
                            float(moyUE2 or 0) * ects_ue2 +
                            float(moyUE3 or 0) * ects_ue3 +
                            float(moyUE4 or 0) * ects_ue4 +
                            float(moyUESPE or 0) * ects_uespe
                        ) / moyenne_ects
                    else:
                        moyenne_ponderee = (
                        float(moyUE1 or 0) * ects_ue1 +
                        float(moyUE2 or 0) * ects_ue2 +
                        float(moyUE3 or 0) * ects_ue3 +
                        float(moyUE4 or 0) * ects_ue4
                    ) / moyenne_ects
                    moyenne_ponderee_str = f"{moyenne_ponderee:.2f}"
                else:
                    moyenne_ponderee_str = ""
            except (ValueError, TypeError, ZeroDivisionError):
                moyenne_ponderee_str = ""


            word_bytes = base64.b64decode(str(word_template.fileData))
            doc = Document(BytesIO(word_bytes))

            # Préparer les données de l'étudiant selon le template
            if template_name == "modeleBG-ALT-S1-2024-2025.docx":
                student_data = {
                    "CodeApprenant": str(updated_ws[f"A{row}"].value or ""),
                    "nomApprenant": str(updated_ws[f"B{row}"].value or ""),
                    "note1": calculate_single_note_average(updated_ws[f"D{row}"].value),
                    "note2": calculate_single_note_average(updated_ws[f"E{row}"].value),
                    "note3": calculate_single_note_average(updated_ws[f"F{row}"].value),
                    "note4": calculate_single_note_average(updated_ws[f"H{row}"].value),
                    "note5": calculate_single_note_average(updated_ws[f"I{row}"].value),
                    "note6": calculate_single_note_average(updated_ws[f"J{row}"].value),
                    "note7": calculate_single_note_average(updated_ws[f"L{row}"].value),
                    "note8": calculate_single_note_average(updated_ws[f"N{row}"].value),
                    "note9": calculate_single_note_average(updated_ws[f"O{row}"].value),
                    "note10": calculate_single_note_average(updated_ws[f"P{row}"].value),
                    "note11": calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    "note12": calculate_single_note_average(updated_ws[f"R{row}"].value),
                    "note13": calculate_single_note_average(updated_ws[f"S{row}"].value),
                    "moyUE1": moyUE1,
                    "moyUE2": moyUE2,
                    "moyUE3": moyUE3,
                    "moyUE4": moyUE4,
                    "moyenne": moyenne_ponderee_str,
                    "dateNaissance": str(updated_ws[f"T{row}"].value or ""),
                    "campus": str(updated_ws[f"U{row}"].value or ""),
                    "groupe": str(updated_ws[f"W{row}"].value or ""),
                    "etendugroupe": str(updated_ws[f"X{row}"].value or ""),
                    "justifiee": str(updated_ws[f"Y{row}"].value or "0h0m"),
                    "injustifiee": str(updated_ws[f"Z{row}"].value or "0h0m"),
                    "retard": str(updated_ws[f"AA{row}"].value or "0h0m"),
                    "APPRECIATIONS": str(updated_ws[f"AB{row}"].value or " "),
                    "datedujour": date_du_jour,
                    "ECTSUE1": str(ects_ue1),
                    "ECTSUE2": str(ects_ue2),
                    "ECTSUE3": str(ects_ue3),
                    "ECTSUE4": str(ects_ue4),
                    "moyenneECTS": str(moyenne_ects),
                    **ue_matieres,
                    **ects_data
                }
                logging.info(f"Date du jour définie dans student_data : {student_data['datedujour']}")


            elif template_name == "modeleBG-ALT-S2-2024-2025.docx":
                student_data = {
                    "CodeApprenant": str(updated_ws[f"A{row}"].value or ""),
                    "nomApprenant": str(updated_ws[f"B{row}"].value or ""),
                    "note1": calculate_single_note_average(updated_ws[f"D{row}"].value),
                    "note2": calculate_single_note_average(updated_ws[f"E{row}"].value),
                    "note3": calculate_single_note_average(updated_ws[f"F{row}"].value),
                    "note4": calculate_single_note_average(updated_ws[f"G{row}"].value),
                    "note5": calculate_single_note_average(updated_ws[f"I{row}"].value),
                    "note6": calculate_single_note_average(updated_ws[f"J{row}"].value),
                    "note7": calculate_single_note_average(updated_ws[f"K{row}"].value),
                    "note8": calculate_single_note_average(updated_ws[f"M{row}"].value),
                    "note9": calculate_single_note_average(updated_ws[f"O{row}"].value),
                    "note10": calculate_single_note_average(updated_ws[f"P{row}"].value),
                    "note11": calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    "note12": calculate_single_note_average(updated_ws[f"R{row}"].value),
                    "note13": calculate_single_note_average(updated_ws[f"S{row}"].value),
                    "moyUE1": moyUE1,
                    "moyUE2": moyUE2,
                    "moyUE3": moyUE3,
                    "moyUE4": moyUE4,
                    "moyenne": moyenne_ponderee_str,
                    "dateNaissance": str(updated_ws[f"T{row}"].value or ""),
                    "campus": str(updated_ws[f"U{row}"].value or ""),
                    "groupe": str(updated_ws[f"V{row}"].value or ""),
                    "etendugroupe": str(updated_ws[f"W{row}"].value or ""),
                    "justifiee": str(updated_ws[f"X{row}"].value or ""),
                    "injustifiee": str(updated_ws[f"Y{row}"].value or ""),
                    "retard": str(updated_ws[f"Z{row}"].value or ""),
                    "APPRECIATIONS": str(updated_ws[f"AA{row}"].value or ""),
                    "datedujour": date_du_jour,
                    "ECTSUE1": str(ects_ue1),
                    "ECTSUE2": str(ects_ue2),
                    "ECTSUE3": str(ects_ue3),
                    "ECTSUE4": str(ects_ue4),
                    "moyenneECTS": str(moyenne_ects),
                    **ue_matieres,
                    **ects_data
                }

            elif template_name == "modeleBG-ALT-S3-2024-2025.docx":
                student_data = {
                    "CodeApprenant": str(updated_ws[f"A{row}"].value or ""),
                    "nomApprenant": str(updated_ws[f"B{row}"].value or ""),
                    "note1": calculate_single_note_average(updated_ws[f"D{row}"].value),
                    "note2": calculate_single_note_average(updated_ws[f"F{row}"].value),
                    "note3": calculate_single_note_average(updated_ws[f"G{row}"].value),
                    "note4": calculate_single_note_average(updated_ws[f"H{row}"].value),
                    "note5": calculate_single_note_average(updated_ws[f"I{row}"].value),
                    "note6": calculate_single_note_average(updated_ws[f"K{row}"].value),
                    "note7": calculate_single_note_average(updated_ws[f"L{row}"].value),
                    "note8": calculate_single_note_average(updated_ws[f"N{row}"].value),
                    "note9": calculate_single_note_average(updated_ws[f"O{row}"].value),
                    "note10": calculate_single_note_average(updated_ws[f"P{row}"].value),
                    "note11": calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    "note12": calculate_single_note_average(updated_ws[f"R{row}"].value),
                    "moyUE1": moyUE1,
                    "moyUE2": moyUE2,
                    "moyUE3": moyUE3,
                    "moyUE4": moyUE4,
                    "moyenne": moyenne_ponderee_str,
                    "dateNaissance": str(updated_ws[f"S{row}"].value or ""),
                    "campus": str(updated_ws[f"T{row}"].value or ""),
                    "groupe": str(updated_ws[f"U{row}"].value or ""),
                    "etendugroupe": str(updated_ws[f"V{row}"].value or ""),
                    "justifiee": str(updated_ws[f"W{row}"].value or ""),
                    "injustifiee": str(updated_ws[f"X{row}"].value or ""),
                    "retard": str(updated_ws[f"Y{row}"].value or ""),
                    "APPRECIATIONS": str(updated_ws[f"Z{row}"].value or ""),
                    "datedujour": date_du_jour,
                    "ECTSUE1": str(ects_ue1),
                    "ECTSUE2": str(ects_ue2),
                    "ECTSUE3": str(ects_ue3),
                    "ECTSUE4": str(ects_ue4),
                    "moyenneECTS": str(moyenne_ects),
                    **ue_matieres,
                    **ects_data
                }

            elif template_name == "modeleBG-ALT-S4-2024-2025.docx":
                student_data = {
                    "CodeApprenant": str(updated_ws[f"A{row}"].value or ""),
                    "nomApprenant": str(updated_ws[f"B{row}"].value or ""),
                    "note1": calculate_single_note_average(updated_ws[f"D{row}"].value),
                    "note2": calculate_single_note_average(updated_ws[f"E{row}"].value),
                    "note3": calculate_single_note_average(updated_ws[f"F{row}"].value),
                    "note4": calculate_single_note_average(updated_ws[f"H{row}"].value),
                    "note5": calculate_single_note_average(updated_ws[f"I{row}"].value),
                    "note6": calculate_single_note_average(updated_ws[f"J{row}"].value),
                    "note7": calculate_single_note_average(updated_ws[f"K{row}"].value),
                    "note8": calculate_single_note_average(updated_ws[f"L{row}"].value),
                    "note9": calculate_single_note_average(updated_ws[f"N{row}"].value),
                    "note10": calculate_single_note_average(updated_ws[f"P{row}"].value),
                    "note11": calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    "note12": calculate_single_note_average(updated_ws[f"R{row}"].value),
                    "note13": calculate_single_note_average(updated_ws[f"S{row}"].value),
                    "moyUE1": moyUE1,
                    "moyUE2": moyUE2,
                    "moyUE3": moyUE3,
                    "moyUE4": moyUE4,
                    "moyenne": moyenne_ponderee_str,
                    "dateNaissance": str(updated_ws[f"T{row}"].value or ""),
                    "campus": str(updated_ws[f"U{row}"].value or ""),
                    "groupe": str(updated_ws[f"V{row}"].value or ""),
                    "etendugroupe": str(updated_ws[f"W{row}"].value or ""),
                    "justifiee": str(updated_ws[f"X{row}"].value or ""),
                    "injustifiee": str(updated_ws[f"Y{row}"].value or ""),
                    "retard": str(updated_ws[f"Z{row}"].value or ""),
                    "APPRECIATIONS": str(updated_ws[f"AA{row}"].value or ""),
                    "datedujour": date_du_jour,
                    "ECTSUE1": str(ects_ue1),
                    "ECTSUE2": str(ects_ue2),
                    "ECTSUE3": str(ects_ue3),
                    "ECTSUE4": str(ects_ue4),
                    "moyenneECTS": str(moyenne_ects),
                    **ue_matieres,
                    **ects_data
                }

            elif template_name == "modeleBG-ALT-S5-2024-2025.docx":
                student_data = {
                    "CodeApprenant": str(updated_ws[f"A{row}"].value or ""),
                    "nomApprenant": str(updated_ws[f"B{row}"].value or ""),
                    "note1": calculate_single_note_average(updated_ws[f"D{row}"].value),
                    "note2": calculate_single_note_average(updated_ws[f"E{row}"].value),
                    "note3": calculate_single_note_average(updated_ws[f"G{row}"].value),
                    "note4": calculate_single_note_average(updated_ws[f"H{row}"].value),
                    "note5": calculate_single_note_average(updated_ws[f"I{row}"].value),
                    "note6": calculate_single_note_average(updated_ws[f"J{row}"].value),
                    "note7": calculate_single_note_average(updated_ws[f"L{row}"].value),
                    "note8": calculate_single_note_average(updated_ws[f"M{row}"].value),
                    "note9": calculate_single_note_average(updated_ws[f"O{row}"].value),
                    "note10": calculate_single_note_average(updated_ws[f"P{row}"].value),
                    "note11": calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    "note12": calculate_single_note_average(updated_ws[f"R{row}"].value),
                    "note13": calculate_single_note_average(updated_ws[f"S{row}"].value),
                    "note14": calculate_single_note_average(updated_ws[f"T{row}"].value),
                    "moyUE1": moyUE1,
                    "moyUE2": moyUE2,
                    "moyUE3": moyUE3,
                    "moyUE4": moyUE4,
                    "moyenne": moyenne_ponderee_str,
                    "dateNaissance": str(updated_ws[f"U{row}"].value or ""),
                    "campus": str(updated_ws[f"V{row}"].value or ""),
                    "groupe": str(updated_ws[f"W{row}"].value or ""),
                    "etendugroupe": str(updated_ws[f"X{row}"].value or ""),
                    "justifiee": str(updated_ws[f"Y{row}"].value or ""),
                    "injustifiee": str(updated_ws[f"Z{row}"].value or ""),
                    "retard": str(updated_ws[f"AA{row}"].value or ""),
                    "APPRECIATIONS": str(updated_ws[f"AB{row}"].value or ""),
                    "datedujour": date_du_jour,
                    "ECTSUE1": str(ects_ue1),
                    "ECTSUE2": str(ects_ue2),
                    "ECTSUE3": str(ects_ue3),
                    "ECTSUE4": str(ects_ue4),
                    "moyenneECTS": str(moyenne_ects),
                    **ue_matieres,
                    **ects_data
                }

            elif template_name == "modeleBG-ALT-S6-2024-2025.docx":
                student_data = {
                    "CodeApprenant": str(updated_ws[f"A{row}"].value or ""),
                    "nomApprenant": str(updated_ws[f"B{row}"].value or ""),
                    "note1": calculate_single_note_average(updated_ws[f"D{row}"].value),
                    "note2": calculate_single_note_average(updated_ws[f"E{row}"].value),
                    "note3": calculate_single_note_average(updated_ws[f"G{row}"].value),
                    "note4": calculate_single_note_average(updated_ws[f"H{row}"].value),
                    "note5": calculate_single_note_average(updated_ws[f"J{row}"].value),
                    "note6": calculate_single_note_average(updated_ws[f"K{row}"].value),
                    "note7": calculate_single_note_average(updated_ws[f"L{row}"].value),
                    "note8": calculate_single_note_average(updated_ws[f"M{row}"].value),
                    "note9": calculate_single_note_average(updated_ws[f"N{row}"].value),
                    "note10": calculate_single_note_average(updated_ws[f"O{row}"].value),
                    "note11": calculate_single_note_average(updated_ws[f"P{row}"].value),
                    "note12": calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    "note13": calculate_single_note_average(updated_ws[f"R{row}"].value),
                    "moyUE1": moyUE1,
                    "moyUE2": moyUE2,
                    "moyUE3": moyUE3,
                    "moyUE4": moyUE4,
                    "moyenne": moyenne_ponderee_str,
                    "dateNaissance": str(updated_ws[f"S{row}"].value or ""),
                    "campus": str(updated_ws[f"T{row}"].value or ""),
                    "groupe": str(updated_ws[f"U{row}"].value or ""),
                    "etendugroupe": str(updated_ws[f"V{row}"].value or ""),
                    "justifiee": str(updated_ws[f"W{row}"].value or ""),
                    "injustifiee": str(updated_ws[f"X{row}"].value or ""),
                    "retard": str(updated_ws[f"Y{row}"].value or ""),
                    "APPRECIATIONS": str(updated_ws[f"Z{row}"].value or ""),
                    "datedujour": date_du_jour,
                    "ECTSUE1": str(ects_ue1),
                    "ECTSUE2": str(ects_ue2),
                    "ECTSUE3": str(ects_ue3),
                    "ECTSUE4": str(ects_ue4),
                    "moyenneECTS": str(moyenne_ects),
                    **ue_matieres,
                    **ects_data
                }
            elif template_name == "modeleBG-TP-S1-2024-2025.docx":
                student_data = {
                    "CodeApprenant": str(updated_ws[f"A{row}"].value or ""),
                    "nomApprenant": str(updated_ws[f"B{row}"].value or ""),
                    "note1": calculate_single_note_average(updated_ws[f"D{row}"].value),
                    "note2": calculate_single_note_average(updated_ws[f"E{row}"].value),
                    "note3": calculate_single_note_average(updated_ws[f"F{row}"].value),
                    "note4": calculate_single_note_average(updated_ws[f"G{row}"].value),
                    "note5": calculate_single_note_average(updated_ws[f"H{row}"].value),
                    "note6": calculate_single_note_average(updated_ws[f"I{row}"].value),
                    "note7": calculate_single_note_average(updated_ws[f"J{row}"].value),
                    "note8": calculate_single_note_average(updated_ws[f"L{row}"].value),
                    "note9": calculate_single_note_average(updated_ws[f"M{row}"].value),
                    "note10": calculate_single_note_average(updated_ws[f"N{row}"].value),
                    "note11": calculate_single_note_average(updated_ws[f"O{row}"].value),
                    "note12": calculate_single_note_average(updated_ws[f"P{row}"].value),
                    "note13": calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    "note14": calculate_single_note_average(updated_ws[f"S{row}"].value),
                    "note15": calculate_single_note_average(updated_ws[f"T{row}"].value),
                    "note16": calculate_single_note_average(updated_ws[f"V{row}"].value),
                    "note17": calculate_single_note_average(updated_ws[f"W{row}"].value),
                    "note18": calculate_single_note_average(updated_ws[f"X{row}"].value),
                    "note19": calculate_single_note_average(updated_ws[f"Y{row}"].value),
                    "note20": calculate_single_note_average(updated_ws[f"Z{row}"].value),
                    "moyUE1": moyUE1,
                    "moyUE2": moyUE2,
                    "moyUE3": moyUE3,
                    "moyUE4": moyUE4,
                    "moyenne": moyenne_ponderee_str,
                    "dateNaissance": str(updated_ws[f"AA{row}"].value or ""),
                    "campus": str(updated_ws[f"AB{row}"].value or ""),
                    "groupe": str(updated_ws[f"AC{row}"].value or ""),
                    "etendugroupe": str(updated_ws[f"AD{row}"].value or ""),
                    "justifiee": str(updated_ws[f"AE{row}"].value or ""),
                    "injustifiee": str(updated_ws[f"AF{row}"].value or ""),
                    "retard": str(updated_ws[f"AG{row}"].value or ""),
                    "APPRECIATIONS": str(updated_ws[f"AH{row}"].value or ""),
                    "datedujour": date_du_jour,
                    "ECTSUE1": str(ects_ue1),
                    "ECTSUE2": str(ects_ue2),
                    "ECTSUE3": str(ects_ue3),
                    "ECTSUE4": str(ects_ue4),
                    "moyenneECTS": str(moyenne_ects),
                    **ue_matieres,
                    **ects_data
                }

            elif template_name == "modeleBG-TP-S2-2024-2025.docx":
                student_data = {
                    "CodeApprenant": str(updated_ws[f"A{row}"].value or ""),
                    "nomApprenant": str(updated_ws[f"B{row}"].value or ""),
                    "note1": calculate_single_note_average(updated_ws[f"D{row}"].value),
                    "note2": calculate_single_note_average(updated_ws[f"E{row}"].value),
                    "moyUE1": moyUE1,
                    "moyenne": moyenne_ponderee_str,
                    "dateNaissance": str(updated_ws[f"F{row}"].value or ""),
                    "campus": str(updated_ws[f"G{row}"].value or ""),
                    "groupe": str(updated_ws[f"H{row}"].value or ""),
                    "etendugroupe": str(updated_ws[f"I{row}"].value or ""),
                    "justifiee": str(updated_ws[f"J{row}"].value or ""),
                    "injustifiee": str(updated_ws[f"K{row}"].value or ""),
                    "retard": str(updated_ws[f"L{row}"].value or ""),
                    "APPRECIATIONS": str(updated_ws[f"M{row}"].value or ""),
                    "datedujour": date_du_jour,
                    "ECTSUE1": str(ects_ue1),
                    "moyenneECTS": str(moyenne_ects),
                    **ue_matieres,
                    **ects_data
                }

            elif template_name == "modeleBG-TP-S3-2024-2025.docx":
                student_data = {
                    "CodeApprenant": str(updated_ws[f"A{row}"].value or ""),
                    "nomApprenant": str(updated_ws[f"B{row}"].value or ""),
                    "note1": calculate_single_note_average(updated_ws[f"D{row}"].value),
                    "note2": calculate_single_note_average(updated_ws[f"F{row}"].value),
                    "note3": calculate_single_note_average(updated_ws[f"G{row}"].value),
                    "note4": calculate_single_note_average(updated_ws[f"H{row}"].value),
                    "note5": calculate_single_note_average(updated_ws[f"I{row}"].value),
                    "note6": calculate_single_note_average(updated_ws[f"J{row}"].value),
                    "note7": calculate_single_note_average(updated_ws[f"K{row}"].value),
                    "note8": calculate_single_note_average(updated_ws[f"L{row}"].value),
                    "note9": calculate_single_note_average(updated_ws[f"M{row}"].value),
                    "note10": calculate_single_note_average(updated_ws[f"O{row}"].value),
                    "note11": calculate_single_note_average(updated_ws[f"P{row}"].value),
                    "note12": calculate_single_note_average(updated_ws[f"R{row}"].value),
                    "note13": calculate_single_note_average(updated_ws[f"S{row}"].value),
                    "note14": calculate_single_note_average(updated_ws[f"T{row}"].value),
                    "note15": calculate_single_note_average(updated_ws[f"U{row}"].value),
                    "moyUE1": moyUE1,
                    "moyUE2": moyUE2,
                    "moyUE3": moyUE3,
                    "moyUE4": moyUE4,
                    "moyenne": moyenne_ponderee_str,
                    "dateNaissance": str(updated_ws[f"V{row}"].value or ""),
                    "campus": str(updated_ws[f"W{row}"].value or ""),
                    "groupe": str(updated_ws[f"X{row}"].value or ""),
                    "etendugroupe": str(updated_ws[f"Y{row}"].value or ""),
                    "justifiee": str(updated_ws[f"Z{row}"].value or ""),
                    "injustifiee": str(updated_ws[f"AA{row}"].value or ""),
                    "retard": str(updated_ws[f"AB{row}"].value or ""),
                    "APPRECIATIONS": str(updated_ws[f"AC{row}"].value or ""),
                    "datedujour": date_du_jour,
                    "ECTSUE1": str(ects_ue1),
                    "ECTSUE2": str(ects_ue2),
                    "ECTSUE3": str(ects_ue3),
                    "ECTSUE4": str(ects_ue4),
                    "moyenneECTS": str(moyenne_ects),
                    **ue_matieres,
                    **ects_data
                }

            elif template_name == "modeleBG-TP-S4-2024-2025.docx":
                student_data = {
                    "CodeApprenant": str(updated_ws[f"A{row}"].value or ""),
                    "nomApprenant": str(updated_ws[f"B{row}"].value or ""),
                    "note1": calculate_single_note_average(updated_ws[f"D{row}"].value),
                    "moyUE1": moyUE1,
                    "moyenne": moyenne_ponderee_str,
                    "dateNaissance": str(updated_ws[f"E{row}"].value or ""),
                    "campus": str(updated_ws[f"F{row}"].value or ""),
                    "groupe": str(updated_ws[f"G{row}"].value or ""),
                    "etendugroupe": str(updated_ws[f"H{row}"].value or ""),
                    "justifiee": str(updated_ws[f"I{row}"].value or ""),
                    "injustifiee": str(updated_ws[f"J{row}"].value or ""),
                    "retard": str(updated_ws[f"K{row}"].value or ""),
                    "APPRECIATIONS": str(updated_ws[f"L{row}"].value or ""),
                    "datedujour": date_du_jour,
                    "ECTSUE1": str(ects_ue1),
                    "moyenneECTS": str(moyenne_ects),
                    **ue_matieres,
                    **ects_data
                }

            elif template_name == "modeleBG-TP-S5-2024-2025.docx":
                student_data = {
                    "CodeApprenant": str(updated_ws[f"A{row}"].value or ""),
                    "nomApprenant": str(updated_ws[f"B{row}"].value or ""),
                    "note1": calculate_single_note_average(updated_ws[f"D{row}"].value),
                    "note2": calculate_single_note_average(updated_ws[f"E{row}"].value),
                    "note3": calculate_single_note_average(updated_ws[f"F{row}"].value),
                    "note4": calculate_single_note_average(updated_ws[f"G{row}"].value),
                    "note5": calculate_single_note_average(updated_ws[f"H{row}"].value),
                    "note6": calculate_single_note_average(updated_ws[f"J{row}"].value),
                    "note7": calculate_single_note_average(updated_ws[f"K{row}"].value),
                    "note8": calculate_single_note_average(updated_ws[f"L{row}"].value),
                    "note9": calculate_single_note_average(updated_ws[f"M{row}"].value),
                    "note10": calculate_single_note_average(updated_ws[f"N{row}"].value),
                    "note11": calculate_single_note_average(updated_ws[f"P{row}"].value),
                    "note12": calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    "note13": calculate_single_note_average(updated_ws[f"R{row}"].value),
                    "note14": calculate_single_note_average(updated_ws[f"S{row}"].value),
                    "note15": calculate_single_note_average(updated_ws[f"T{row}"].value),
                    "note16": calculate_single_note_average(updated_ws[f"V{row}"].value),
                    "note17": calculate_single_note_average(updated_ws[f"W{row}"].value),
                    "note18": calculate_single_note_average(updated_ws[f"X{row}"].value),
                    "note19": calculate_single_note_average(updated_ws[f"Y{row}"].value),
                    "note20": calculate_single_note_average(updated_ws[f"Z{row}"].value),
                    "moyUE1": moyUE1,
                    "moyUE2": moyUE2,
                    "moyUE3": moyUE3,
                    "moyUE4": moyUE4,
                    "moyenne": moyenne_ponderee_str,
                    "dateNaissance": str(updated_ws[f"AA{row}"].value or ""),
                    "campus": str(updated_ws[f"AB{row}"].value or ""),
                    "groupe": str(updated_ws[f"AC{row}"].value or ""),
                    "etendugroupe": str(updated_ws[f"AD{row}"].value or ""),
                    "justifiee": str(updated_ws[f"AE{row}"].value or ""),
                    "injustifiee": str(updated_ws[f"AF{row}"].value or ""),
                    "retard": str(updated_ws[f"AG{row}"].value or ""),
                    "APPRECIATIONS": str(updated_ws[f"AH{row}"].value or ""),
                    "datedujour": date_du_jour,
                    "ECTSUE1": str(ects_ue1),
                    "ECTSUE2": str(ects_ue2),
                    "ECTSUE3": str(ects_ue3),
                    "ECTSUE4": str(ects_ue4),
                    "moyenneECTS": str(moyenne_ects),
                    **ue_matieres,
                    **ects_data
                }

            elif template_name == "modeleBG-TP-S6-2024-2025.docx":
                student_data = {
                    "CodeApprenant": str(updated_ws[f"A{row}"].value or ""),
                    "nomApprenant": str(updated_ws[f"B{row}"].value or ""),
                    "note1": calculate_single_note_average(updated_ws[f"D{row}"].value),
                    "note2": calculate_single_note_average(updated_ws[f"E{row}"].value),
                    "note3": calculate_single_note_average(updated_ws[f"F{row}"].value),
                    "moyUE1": moyUE1,
                    "moyenne": moyenne_ponderee_str,
                    "dateNaissance": str(updated_ws[f"G{row}"].value or ""),
                    "campus": str(updated_ws[f"H{row}"].value or ""),
                    "groupe": str(updated_ws[f"I{row}"].value or ""),
                    "etendugroupe": str(updated_ws[f"J{row}"].value or ""),
                    "justifiee": str(updated_ws[f"K{row}"].value or ""),
                    "injustifiee": str(updated_ws[f"L{row}"].value or ""),
                    "retard": str(updated_ws[f"M{row}"].value or ""),
                    "APPRECIATIONS": str(updated_ws[f"N{row}"].value or ""),
                    "datedujour": date_du_jour,
                    "ECTSUE1": str(ects_ue1),
                    "moyenneECTS": str(moyenne_ects),
                    **ue_matieres,
                    **ects_data
                }

            elif template_name == "modeleM1-S1.docx":
                student_data = {
                    "CodeApprenant": str(updated_ws[f"A{row}"].value or ""),
                    "nomApprenant": str(updated_ws[f"B{row}"].value or ""),
                    "note1": calculate_single_note_average(updated_ws[f"D{row}"].value),
                    "note2": calculate_single_note_average(updated_ws[f"E{row}"].value),
                    "note3": calculate_single_note_average(updated_ws[f"G{row}"].value),
                    "note4": calculate_single_note_average(updated_ws[f"H{row}"].value),
                    "note5": calculate_single_note_average(updated_ws[f"J{row}"].value),
                    "note6": calculate_single_note_average(updated_ws[f"K{row}"].value),
                    "note7": calculate_single_note_average(updated_ws[f"M{row}"].value),
                    "note8": calculate_single_note_average(updated_ws[f"O{row}"].value),
                    "note9": calculate_single_note_average(updated_ws[f"P{row}"].value),
                    "note10": calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    "note11": calculate_single_note_average(updated_ws[f"R{row}"].value),
                    "note12": calculate_single_note_average(updated_ws[f"T{row}"].value),
                    "note13": calculate_single_note_average(updated_ws[f"U{row}"].value),
                    "note14": calculate_single_note_average(updated_ws[f"V{row}"].value),
                    "moyUE1": moyUE1,
                    "moyUE2": moyUE2,
                    "moyUE3": moyUE3,
                    "moyUE4": moyUE4,
                    "moyUESPE": moyUESPE,
                    "moyenne": moyenne_ponderee_str,
                    "dateNaissance": str(updated_ws[f"W{row}"].value or ""),
                    "campus": str(updated_ws[f"X{row}"].value or ""),
                    "groupe": str(updated_ws[f"Y{row}"].value or ""),
                    "etendugroupe": str(updated_ws[f"Z{row}"].value or ""),
                    "justifiee": str(updated_ws[f"AA{row}"].value or ""),
                    "injustifiee": str(updated_ws[f"AB{row}"].value or ""),
                    "retard": str(updated_ws[f"AC{row}"].value or ""),
                    "APPRECIATIONS": str(updated_ws[f"AD{row}"].value or ""),
                    "datedujour": date_du_jour,
                    "ECTSUE1": str(ects_ue1),
                    "ECTSUE2": str(ects_ue2),
                    "ECTSUE3": str(ects_ue3),
                    "ECTSUE4": str(ects_ue4),
                    "ECTSUESPE": str(ects_uespe),
                    "moyenneECTS": str(moyenne_ects),
                    **ue_matieres,
                    **ects_data
                }

            elif template_name == "modeleM2-S3.docx":
                student_data = {
                    "CodeApprenant": str(updated_ws[f"A{row}"].value or ""),
                    "nomApprenant": str(updated_ws[f"B{row}"].value or ""),
                    "note1": calculate_single_note_average(updated_ws[f"D{row}"].value),
                    "note2": calculate_single_note_average(updated_ws[f"E{row}"].value),
                    "note3": calculate_single_note_average(updated_ws[f"G{row}"].value),
                    "note4": calculate_single_note_average(updated_ws[f"I{row}"].value),
                    "note5": calculate_single_note_average(updated_ws[f"K{row}"].value),
                    "note6": calculate_single_note_average(updated_ws[f"L{row}"].value),
                    "note7": calculate_single_note_average(updated_ws[f"M{row}"].value),
                    "note8": calculate_single_note_average(updated_ws[f"N{row}"].value),
                    "note9": calculate_single_note_average(updated_ws[f"O{row}"].value),
                    "note10": calculate_single_note_average(updated_ws[f"Q{row}"].value),
                    "note11": calculate_single_note_average(updated_ws[f"R{row}"].value),
                    "note12": calculate_single_note_average(updated_ws[f"S{row}"].value),
                    "note13": calculate_single_note_average(updated_ws[f"T{row}"].value),
                    "moyUE1": moyUE1,
                    "moyUE2": moyUE2,
                    "moyUE3": moyUE3,
                    "moyUE4": moyUE4,
                    "moyUESPE": moyUESPE,
                    "moyenne": moyenne_ponderee_str,
                    "dateNaissance": str(updated_ws[f"U{row}"].value or ""),
                    "campus": str(updated_ws[f"V{row}"].value or ""),
                    "groupe": str(updated_ws[f"W{row}"].value or ""),
                    "etendugroupe": str(updated_ws[f"X{row}"].value or ""),
                    "justifiee": str(updated_ws[f"Y{row}"].value or ""),
                    "injustifiee": str(updated_ws[f"Z{row}"].value or ""),
                    "retard": str(updated_ws[f"AA{row}"].value or ""),
                    "APPRECIATIONS": str(updated_ws[f"AB{row}"].value or ""),
                    "datedujour": date_du_jour,
                    "ECTSUE1": str(ects_ue1),
                    "ECTSUE2": str(ects_ue2),
                    "ECTSUE3": str(ects_ue3),
                    "ECTSUE4": str(ects_ue4),
                    "ECTSUESPE": str(ects_uespe),
                    "moyenneECTS": str(moyenne_ects),
                    **ue_matieres,
                    **ects_data
                }


            # Calculer les états pour chaque note
            # Calculer les états pour chaque note
            etats = {}
            if template_name == "modeleBG-ALT-S1-2024-2025.docx":
                # Vérifier d'abord s'il y a des R dans chaque UE
                has_r_ue1 = any(float(student_data[f"note{i}"]) < 8 for i in range(1, 4) if student_data[f"note{i}"])
                has_r_ue2 = any(float(student_data[f"note{i}"]) < 8 for i in range(4, 7) if student_data[f"note{i}"])
                has_r_ue3 = float(student_data["note7"]) < 8 if student_data["note7"] else False
                has_r_ue4 = any(float(student_data[f"note{i}"]) < 8 for i in range(8, 14) if student_data[f"note{i}"])

                # Calculer les états en tenant compte des R
                # UE1 (3 notes)
                for i in range(1, 4):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue1)
                
                # UE2 (3 notes)
                for i in range(4, 7):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue2)
                
                # UE3 (1 note)
                etats["etat7"] = get_etat(student_data["note7"], has_r_ue3)
                
                # UE4 (6 notes)
                for i in range(8, 14):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue4)



            elif template_name == "modeleBG-ALT-S2-2024-2025.docx":
                has_r_ue1 = any(float(student_data[f"note{i}"]) < 8 for i in range(1, 5) if student_data[f"note{i}"])
                has_r_ue2 = any(float(student_data[f"note{i}"]) < 8 for i in range(5, 8) if student_data[f"note{i}"])
                has_r_ue3 = float(student_data["note8"]) < 8 if student_data["note8"] else False
                has_r_ue4 = any(float(student_data[f"note{i}"]) < 8 for i in range(9, 16) if student_data[f"note{i}"])

                for i in range(1, 5):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue1)
                for i in range(5, 8):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue2)
                etats["etat8"] = get_etat(student_data["note8"], has_r_ue3)
                for i in range(9, 16):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue4)

            elif template_name == "modeleBG-ALT-S3-2024-2025.docx":
                has_r_ue1 = any(float(student_data[f"note{i}"]) < 8 for i in range(1, 6) if student_data[f"note{i}"])
                has_r_ue2 = any(float(student_data[f"note{i}"]) < 8 for i in range(6, 8) if student_data[f"note{i}"])
                has_r_ue3 = float(student_data["note8"]) < 8 if student_data["note8"] else False
                has_r_ue4 = any(float(student_data[f"note{i}"]) < 8 for i in range(9, 14) if student_data[f"note{i}"])

                for i in range(1, 6):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue1)
                for i in range(6, 8):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue2)
                etats["etat8"] = get_etat(student_data["note8"], has_r_ue3)
                for i in range(9, 14):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue4)

            elif template_name == "modeleBG-ALT-S4-2024-2025.docx":
                has_r_ue1 = any(float(student_data[f"note{i}"]) < 8 for i in range(1, 5) if student_data[f"note{i}"])
                has_r_ue2 = any(float(student_data[f"note{i}"]) < 8 for i in range(5, 8) if student_data[f"note{i}"])
                has_r_ue3 = float(student_data["note8"]) < 8 if student_data["note8"] else False
                has_r_ue4 = any(float(student_data[f"note{i}"]) < 8 for i in range(9, 14) if student_data[f"note{i}"])

                for i in range(1, 5):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue1)
                for i in range(5, 8):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue2)
                etats["etat8"] = get_etat(student_data["note8"], has_r_ue3)
                for i in range(9, 14):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue4)

            elif template_name == "modeleBG-ALT-S5-2024-2025.docx":
                has_r_ue1 = any(float(student_data[f"note{i}"]) < 8 for i in range(1, 4) if student_data[f"note{i}"])
                has_r_ue2 = any(float(student_data[f"note{i}"]) < 8 for i in range(4, 7) if student_data[f"note{i}"])
                has_r_ue3 = float(student_data["note7"]) < 8 if student_data["note7"] else False
                has_r_ue4 = any(float(student_data[f"note{i}"]) < 8 for i in range(8, 15) if student_data[f"note{i}"])

                for i in range(1, 4):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue1)
                for i in range(4, 7):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue2)
                etats["etat7"] = get_etat(student_data["note7"], has_r_ue3)
                for i in range(8, 15):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue4)

            elif template_name == "modeleBG-ALT-S6-2024-2025.docx":
                has_r_ue1 = any(float(student_data[f"note{i}"]) < 8 for i in range(1, 4) if student_data[f"note{i}"])
                has_r_ue2 = any(float(student_data[f"note{i}"]) < 8 for i in range(4, 7) if student_data[f"note{i}"])
                has_r_ue3 = float(student_data["note7"]) < 8 if student_data["note7"] else False
                has_r_ue4 = any(float(student_data[f"note{i}"]) < 8 for i in range(8, 13) if student_data[f"note{i}"])

                for i in range(1, 4):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue1)
                for i in range(4, 7):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue2)
                etats["etat7"] = get_etat(student_data["note7"], has_r_ue3)
                for i in range(8, 13):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue4)
            
            elif template_name == "modeleBG-TP-S1-2024-2025.docx":
                # Vérifier d'abord s'il y a des R dans chaque UE
                has_r_ue1 = any(float(student_data[f"note{i}"]) < 8 for i in range(1, 6) if student_data[f"note{i}"])
                has_r_ue2 = any(float(student_data[f"note{i}"]) < 8 for i in range(6, 9) if student_data[f"note{i}"])
                has_r_ue3 = any(float(student_data[f"note{i}"]) < 8 for i in range(9, 12) if student_data[f"note{i}"])
                has_r_ue4 = any(float(student_data[f"note{i}"]) < 8 for i in range(12, 22) if student_data[f"note{i}"])

                # Calculer les états en tenant compte des R
                for i in range(1, 6):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue1)
                for i in range(6, 9):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue2)
                for i in range(9, 12):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue3)
                for i in range(12, 22):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue4)

            elif template_name == "modeleBG-TP-S2-2024-2025.docx":
                has_r_ue1 = float(student_data["note1"]) < 8 if student_data["note1"] else False
                has_r_ue2 = float(student_data["note2"]) < 8 if student_data["note2"] else False

                etats["etat1"] = get_etat(student_data["note1"], has_r_ue1)
                etats["etat2"] = get_etat(student_data["note2"], has_r_ue2)

            elif template_name == "modeleBG-TP-S3-2024-2025.docx":
                has_r_ue1 = any(float(student_data[f"note{i}"]) < 8 for i in range(1, 7) if student_data[f"note{i}"])
                has_r_ue2 = any(float(student_data[f"note{i}"]) < 8 for i in range(7, 10) if student_data[f"note{i}"])
                has_r_ue3 = any(float(student_data[f"note{i}"]) < 8 for i in range(10, 13) if student_data[f"note{i}"])
                has_r_ue4 = any(float(student_data[f"note{i}"]) < 8 for i in range(13, 15) if student_data[f"note{i}"])

                for i in range(1, 7):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue1)
                for i in range(7, 10):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue2)
                for i in range(10, 13):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue3)
                for i in range(13, 15):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue4)

            elif template_name == "modeleBG-TP-S4-2024-2025.docx":
                has_r_ue1 = float(student_data["note1"]) < 8 if student_data["note1"] else False

                etats["etat1"] = get_etat(student_data["note1"], has_r_ue1)

            elif template_name == "modeleBG-TP-S5-2024-2025.docx":
                has_r_ue1 = any(float(student_data[f"note{i}"]) < 8 for i in range(1, 7) if student_data[f"note{i}"])
                has_r_ue2 = any(float(student_data[f"note{i}"]) < 8 for i in range(7, 11) if student_data[f"note{i}"])
                has_r_ue3 = any(float(student_data[f"note{i}"]) < 8 for i in range(11, 15) if student_data[f"note{i}"])
                has_r_ue4 = any(float(student_data[f"note{i}"]) < 8 for i in range(15, 19) if student_data[f"note{i}"])

                for i in range(1, 7):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue1)
                for i in range(7, 11):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue2)
                for i in range(11, 15):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue3)
                for i in range(15, 19):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue4)

            elif template_name == "modeleBG-TP-S6-2024-2025.docx":
                has_r_ue1 = float(student_data["note1"]) < 8 if student_data["note1"] else False
                has_r_ue2 = float(student_data["note2"]) < 8 if student_data["note2"] else False
                has_r_ue3 = float(student_data["note3"]) < 8 if student_data["note3"] else False

                etats["etat1"] = get_etat(student_data["note1"], has_r_ue1)
                etats["etat2"] = get_etat(student_data["note2"], has_r_ue2)
                etats["etat3"] = get_etat(student_data["note3"], has_r_ue3)

            elif template_name == "modeleM1-S1.docx":
                has_r_ue1 = any(float(student_data[f"note{i}"]) < 8 for i in range(1, 4) if student_data[f"note{i}"])
                has_r_ue2 = any(float(student_data[f"note{i}"]) < 8 for i in range(4, 6) if student_data[f"note{i}"])
                has_r_ue3 = float(student_data["note6"]) < 8 if student_data["note6"] else False
                has_r_ue4 = any(float(student_data[f"note{i}"]) < 8 for i in range(7, 15) if student_data[f"note{i}"])

                for i in range(1, 4):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue1)
                for i in range(4, 6):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue2)
                etats["etat6"] = get_etat(student_data["note6"], has_r_ue3)
                for i in range(7, 15):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue4)

            elif template_name == "modeleM2-S3.docx":
                has_r_ue1 = any(float(student_data[f"note{i}"]) < 8 for i in range(1, 6) if student_data[f"note{i}"])
                has_r_ue2 = any(float(student_data[f"note{i}"]) < 8 for i in range(6, 8) if student_data[f"note{i}"])
                has_r_ue3 = float(student_data["note8"]) < 8 if student_data["note8"] else False
                has_r_ue4 = any(float(student_data[f"note{i}"]) < 8 for i in range(9, 14) if student_data[f"note{i}"])

                for i in range(1, 6):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue1)
                for i in range(6, 8):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue2)
                etats["etat8"] = get_etat(student_data["note8"], has_r_ue3)
                for i in range(9, 14):
                    etats[f"etat{i}"] = get_etat(student_data[f"note{i}"], has_r_ue4)

            student_data.update(etats)


            # Calculer les états des UE
            if template_name == "modeleBG-ALT-S1-2024-2025.docx":
                etats_ue = {
                    "etatUE1": get_etat_ue([etats[f"etat{i}"] for i in range(1, 4)], student_data["moyUE1"]),  # 3 notes
                    "etatUE2": get_etat_ue([etats[f"etat{i}"] for i in range(4, 7)], student_data["moyUE2"]),  # 3 notes
                    "etatUE3": get_etat_ue([etats["etat7"]], student_data["moyUE3"]),  # 1 note
                    "etatUE4": get_etat_ue([etats[f"etat{i}"] for i in range(8, 14)], student_data["moyUE4"])  # 6 notes
                }



            elif template_name == "modeleBG-ALT-S2-2024-2025.docx":
                etats_ue = {
                    "etatUE1": get_etat_ue([etats[f"etat{i}"] for i in range(1, 5)], student_data["moyUE1"]),
                    "etatUE2": get_etat_ue([etats[f"etat{i}"] for i in range(5, 8)], student_data["moyUE2"]),
                    "etatUE3": get_etat_ue([etats["etat8"]], student_data["moyUE3"]),
                    "etatUE4": get_etat_ue([etats[f"etat{i}"] for i in range(9, 16)], student_data["moyUE4"])
                }
            elif template_name == "modeleBG-ALT-S3-2024-2025.docx":
                etats_ue = {
                    "etatUE1": get_etat_ue([etats[f"etat{i}"] for i in range(1, 6)], student_data["moyUE1"]),
                    "etatUE2": get_etat_ue([etats[f"etat{i}"] for i in range(6, 8)], student_data["moyUE2"]),
                    "etatUE3": get_etat_ue([etats["etat8"]], student_data["moyUE3"]),
                    "etatUE4": get_etat_ue([etats[f"etat{i}"] for i in range(9, 14)], student_data["moyUE4"])
                }
            elif template_name == "modeleBG-ALT-S4-2024-2025.docx":
                etats_ue = {
                    "etatUE1": get_etat_ue([etats[f"etat{i}"] for i in range(1, 5)], student_data["moyUE1"]),
                    "etatUE2": get_etat_ue([etats[f"etat{i}"] for i in range(5, 8)], student_data["moyUE2"]),
                    "etatUE3": get_etat_ue([etats["etat8"]], student_data["moyUE3"]),
                    "etatUE4": get_etat_ue([etats[f"etat{i}"] for i in range(9, 14)], student_data["moyUE4"])
                }
            elif template_name == "modeleBG-ALT-S5-2024-2025.docx":
                etats_ue = {
                    "etatUE1": get_etat_ue([etats[f"etat{i}"] for i in range(1, 6)], student_data["moyUE1"]),
                    "etatUE2": get_etat_ue([etats[f"etat{i}"] for i in range(6, 8)], student_data["moyUE2"]),
                    "etatUE3": get_etat_ue([etats["etat8"]], student_data["moyUE3"]),
                    "etatUE4": get_etat_ue([etats[f"etat{i}"] for i in range(9, 15)], student_data["moyUE4"])
                }
            elif template_name == "modeleBG-ALT-S6-2024-2025.docx":
                etats_ue = {
                    "etatUE1": get_etat_ue([etats[f"etat{i}"] for i in range(1, 4)], student_data["moyUE1"]),
                    "etatUE2": get_etat_ue([etats[f"etat{i}"] for i in range(4, 6)], student_data["moyUE2"]),
                    "etatUE3": get_etat_ue([etats["etat6"]], student_data["moyUE3"]),
                    "etatUE4": get_etat_ue([etats[f"etat{i}"] for i in range(7, 13)], student_data["moyUE4"])
                }
            elif template_name == "modeleBG-TP-S1-2024-2025.docx":
                etats_ue = {
                    "etatUE1": get_etat_ue([etats[f"etat{i}"] for i in range(1, 6)], student_data["moyUE1"]),
                    "etatUE2": get_etat_ue([etats[f"etat{i}"] for i in range(6, 9)], student_data["moyUE2"]),
                    "etatUE3": get_etat_ue([etats[f"etat{i}"] for i in range(9, 12)], student_data["moyUE3"]),
                    "etatUE4": get_etat_ue([etats[f"etat{i}"] for i in range(12, 22)], student_data["moyUE4"])
                }
            elif template_name == "modeleBG-TP-S2-2024-2025.docx":
                etats_ue = {
                    "etatUE1": get_etat_ue([etats["etat1"]], student_data["moyUE1"]),
                    "etatUE2": get_etat_ue([etats["etat2"]], student_data["moyUE2"])
                }
            elif template_name == "modeleBG-TP-S3-2024-2025.docx":
                etats_ue = {
                    "etatUE1": get_etat_ue([etats[f"etat{i}"] for i in range(1, 7)], student_data["moyUE1"]),
                    "etatUE2": get_etat_ue([etats[f"etat{i}"] for i in range(7, 10)], student_data["moyUE2"]),
                    "etatUE3": get_etat_ue([etats[f"etat{i}"] for i in range(10, 13)], student_data["moyUE3"]),
                    "etatUE4": get_etat_ue([etats[f"etat{i}"] for i in range(13, 15)], student_data["moyUE4"])
                }
            elif template_name == "modeleBG-TP-S4-2024-2025.docx":
                etats_ue = {
                    "etatUE1": get_etat_ue([etats["etat1"]], student_data["moyUE1"])
                }
            elif template_name == "modeleBG-TP-S5-2024-2025.docx":
                etats_ue = {
                    "etatUE1": get_etat_ue([etats[f"etat{i}"] for i in range(1, 7)], student_data["moyUE1"]),
                    "etatUE2": get_etat_ue([etats[f"etat{i}"] for i in range(7, 11)], student_data["moyUE2"]),
                    "etatUE3": get_etat_ue([etats[f"etat{i}"] for i in range(11, 15)], student_data["moyUE3"]),
                    "etatUE4": get_etat_ue([etats[f"etat{i}"] for i in range(15, 19)], student_data["moyUE4"])
                }
            elif template_name == "modeleBG-TP-S6-2024-2025.docx":
                etats_ue = {
                    "etatUE1": get_etat_ue([etats["etat1"]], student_data["moyUE1"]),
                    "etatUE2": get_etat_ue([etats["etat2"]], student_data["moyUE2"]),
                    "etatUE3": get_etat_ue([etats["etat3"]], student_data["moyUE3"])
                }
            elif template_name == "modeleM1-S1.docx":
                etats_ue = {
                    "etatUE1": get_etat_ue([etats[f"etat{i}"] for i in range(1, 4)], student_data["moyUE1"]),
                    "etatUE2": get_etat_ue([etats[f"etat{i}"] for i in range(4, 6)], student_data["moyUE2"]),
                    "etatUE3": get_etat_ue([etats["etat6"]], student_data["moyUE3"]),
                    "etatUE4": get_etat_ue([etats[f"etat{i}"] for i in range(7, 15)], student_data["moyUE4"])
                }
            elif template_name == "modeleM2-S3.docx":
                etats_ue = {
                    "etatUE1": get_etat_ue([etats[f"etat{i}"] for i in range(1, 6)], student_data["moyUE1"]),
                    "etatUE2": get_etat_ue([etats[f"etat{i}"] for i in range(6, 8)], student_data["moyUE2"]),
                    "etatUE3": get_etat_ue([etats["etat8"]], student_data["moyUE3"]),
                    "etatUE4": get_etat_ue([etats[f"etat{i}"] for i in range(9, 14)], student_data["moyUE4"])
                }

            student_data.update(etats_ue)

            # Calculer le total des états
            if template_name == "modeleBG-TP-S2-2024-2025.docx":
                student_data["totaletat"] = get_total_etat(
                    etats_ue["etatUE1"],
                    etats_ue["etatUE2"]
                )
            elif template_name == "modeleBG-TP-S4-2024-2025.docx":
                student_data["totaletat"] = etats_ue["etatUE1"]
            elif template_name == "modeleBG-TP-S6-2024-2025.docx":
                student_data["totaletat"] = get_total_etat(
                    etats_ue["etatUE1"],
                    etats_ue["etatUE2"],
                    etats_ue["etatUE3"]
                )
            else:
                student_data["totaletat"] = get_total_etat(
                    etats_ue["etatUE1"],
                    etats_ue["etatUE2"],
                    etats_ue["etatUE3"],
                    etats_ue["etatUE4"]
                )



            # Ajuster les ECTS en fonction des notes
            # Ajuster les ECTS en fonction des notes
            
            def adjust_ects(note_str, original_ects, is_ue_moyenne=False, etat=""):
                """
                Ajuste les ECTS en fonction de la note et de son état
                - Pour les moyennes d'UE (is_ue_moyenne=True): garde toujours l'ECTS original
                - Pour les notes individuelles: met l'ECTS à 0 si la note est < 8 ou si l'état est "R"
                """
                if is_ue_moyenne:
                    return str(original_ects)
                try:
                    note = float(note_str) if note_str else 0
                    # Mettre l'ECTS à 0 si la note est < 8 OU si l'état est "R"
                    return "0" if (note < 8 or etat == "R") else str(original_ects)
                except (ValueError, TypeError):
                    return str(original_ects)

            # Ajuster les ECTS pour chaque matière
            if template_name == "modeleBG-ALT-S1-2024-2025.docx":
                # ECTS pour les notes individuelles
                for i in range(1, 14):
                    student_data[f"ECTS{i}"] = adjust_ects(student_data[f"note{i}"], ects_data[f"ECTS{i}"])
                
                # ECTS pour les moyennes d'UE
                student_data["ECTSUE1"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(1, 4)))
                student_data["ECTSUE2"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(4, 7)))
                student_data["ECTSUE3"] = student_data["ECTS7"]
                student_data["ECTSUE4"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(8, 14)))

            elif template_name == "modeleBG-ALT-S2-2024-2025.docx":
                # ECTS pour les notes individuelles
                for i in range(1, 16):
                    student_data[f"ECTS{i}"] = adjust_ects(student_data[f"note{i}"], ects_data[f"ECTS{i}"])
                
                # ECTS pour les moyennes d'UE
                student_data["ECTSUE1"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(1, 5)))
                student_data["ECTSUE2"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(5, 8)))
                student_data["ECTSUE3"] = student_data["ECTS8"]
                student_data["ECTSUE4"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(9, 16)))

            elif template_name == "modeleBG-ALT-S3-2024-2025.docx":
                # ECTS pour les notes individuelles
                for i in range(1, 14):
                    student_data[f"ECTS{i}"] = adjust_ects(student_data[f"note{i}"], ects_data[f"ECTS{i}"])
                
                # ECTS pour les moyennes d'UE
                student_data["ECTSUE1"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(1, 6)))
                student_data["ECTSUE2"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(6, 8)))
                student_data["ECTSUE3"] = student_data["ECTS8"]
                student_data["ECTSUE4"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(9, 14)))

            elif template_name == "modeleBG-ALT-S4-2024-2025.docx":
                # ECTS pour les notes individuelles
                for i in range(1, 14):
                    student_data[f"ECTS{i}"] = adjust_ects(student_data[f"note{i}"], ects_data[f"ECTS{i}"])
                
                # ECTS pour les moyennes d'UE
                student_data["ECTSUE1"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(1, 5)))
                student_data["ECTSUE2"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(5, 8)))
                student_data["ECTSUE3"] = student_data["ECTS8"]
                student_data["ECTSUE4"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(9, 14)))

            elif template_name == "modeleBG-ALT-S5-2024-2025.docx":
                # ECTS pour les notes individuelles
                for i in range(1, 15):
                    student_data[f"ECTS{i}"] = adjust_ects(student_data[f"note{i}"], ects_data[f"ECTS{i}"])
                
                # ECTS pour les moyennes d'UE
                student_data["ECTSUE1"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(1, 4)))
                student_data["ECTSUE2"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(4, 7)))
                student_data["ECTSUE3"] = student_data["ECTS7"]
                student_data["ECTSUE4"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(8, 15)))

            elif template_name == "modeleBG-ALT-S6-2024-2025.docx":
                # ECTS pour les notes individuelles
                for i in range(1, 13):
                    student_data[f"ECTS{i}"] = adjust_ects(student_data[f"note{i}"], ects_data[f"ECTS{i}"])
                
                # ECTS pour les moyennes d'UE
                student_data["ECTSUE1"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(1, 4)))
                student_data["ECTSUE2"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(4, 7)))
                student_data["ECTSUE3"] = student_data["ECTS7"]
                student_data["ECTSUE4"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(8, 13)))
            elif template_name == "modeleBG-TP-S1-2024-2025.docx":
                # ECTS pour les notes individuelles
                for i in range(1, 21):
                    student_data[f"ECTS{i}"] = adjust_ects(student_data[f"note{i}"], ects_data[f"ECTS{i}"])
                
                # ECTS pour les moyennes d'UE
                student_data["ECTSUE1"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(1, 8)))
                student_data["ECTSUE2"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(8, 14)))
                student_data["ECTSUE3"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(14, 16)))
                student_data["ECTSUE4"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(16, 21)))

            elif template_name == "modeleBG-TP-S2-2024-2025.docx":
                # ECTS pour les notes individuelles
                for i in range(1, 3):
                    student_data[f"ECTS{i}"] = adjust_ects(student_data[f"note{i}"], ects_data[f"ECTS{i}"])
                
                # ECTS pour les moyennes d'UE
                student_data["ECTSUE1"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(1, 3)))

            elif template_name == "modeleBG-TP-S3-2024-2025.docx":
                # ECTS pour les notes individuelles
                for i in range(1, 16):
                    student_data[f"ECTS{i}"] = adjust_ects(student_data[f"note{i}"], ects_data[f"ECTS{i}"])
                
                # ECTS pour les moyennes d'UE
                student_data["ECTSUE1"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(1, 2)))
                student_data["ECTSUE2"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(2, 10)))
                student_data["ECTSUE3"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(10, 12)))
                student_data["ECTSUE4"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(12, 16)))

            elif template_name == "modeleBG-TP-S4-2024-2025.docx":
                # ECTS pour les notes individuelles
                student_data["ECTS1"] = adjust_ects(student_data["note1"], ects_data["ECTS1"])
                
                # ECTS pour les moyennes d'UE
                student_data["ECTSUE1"] = student_data["ECTS1"]

            elif template_name == "modeleBG-TP-S5-2024-2025.docx":
                # ECTS pour les notes individuelles
                for i in range(1, 21):
                    student_data[f"ECTS{i}"] = adjust_ects(student_data[f"note{i}"], ects_data[f"ECTS{i}"])
                
                # ECTS pour les moyennes d'UE
                student_data["ECTSUE1"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(1, 6)))
                student_data["ECTSUE2"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(6, 11)))
                student_data["ECTSUE3"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(11, 16)))
                student_data["ECTSUE4"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(16, 21)))

            elif template_name == "modeleBG-TP-S6-2024-2025.docx":
                # ECTS pour les notes individuelles
                for i in range(1, 4):
                    student_data[f"ECTS{i}"] = adjust_ects(student_data[f"note{i}"], ects_data[f"ECTS{i}"])
                
                # ECTS pour les moyennes d'UE
                student_data["ECTSUE1"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(1, 4)))

            elif template_name == "modeleM1-S1.docx":
                # ECTS pour les notes individuelles
                for i in range(1, 15):
                    student_data[f"ECTS{i}"] = adjust_ects(student_data[f"note{i}"], ects_data[f"ECTS{i}"])
                
                # ECTS pour les moyennes d'UE
                student_data["ECTSUE1"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(1, 3)))
                student_data["ECTSUE2"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(3, 5)))
                student_data["ECTSUE3"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(5, 7)))
                student_data["ECTSUE4"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(7, 12)))
                student_data["ECTSUESPE"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(12, 15)))

            elif template_name == "modeleM2-S3.docx":
                # ECTS pour les notes individuelles
                for i in range(1, 14):
                    student_data[f"ECTS{i}"] = adjust_ects(student_data[f"note{i}"], ects_data[f"ECTS{i}"])
                
                # ECTS pour les moyennes d'UE
                student_data["ECTSUE1"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(1, 3)))
                student_data["ECTSUE2"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(3, 4)))
                student_data["ECTSUE3"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(4, 5)))
                student_data["ECTSUE4"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(5, 10)))
                student_data["ECTSUESPE"] = str(sum(int(student_data[f"ECTS{i}"]) for i in range(10, 14)))

            # Calculer le total des ECTS pour tous les templates
            if "ECTSUESPE" in student_data:
                student_data["moyenneECTS"] = str(sum(int(student_data[f"ECTSUE{i}"]) for i in range(1, 5)) + int(student_data["ECTSUESPE"]))
            else:
                student_data["moyenneECTS"] = str(sum(int(student_data[f"ECTSUE{i}"]) for i in range(1, 5)))


            # Remplacer les variables dans le document
            # Remplacer les variables dans le document
            # Remplacer les variables dans le document
            for paragraph in doc.paragraphs:
                for key, value in student_data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in paragraph.text:
                        if key == "etendugroupe":
                            # Sauvegarder le texte original
                            original_text = paragraph.text
                            # Trouver la position du placeholder
                            placeholder_start = original_text.find(placeholder)
                            # Obtenir le texte après le placeholder
                            text_after = original_text[placeholder_start + len(placeholder):]
                            # Effacer le texte du paragraphe
                            paragraph.text = ""
                            # Ajouter la valeur avec le style voulu
                            run = paragraph.add_run(str(value))
                            run.bold = True
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)
                            run.font.name = 'Poppins'
                            # Ajouter le texte qui suivait le placeholder
                            if text_after:
                                run = paragraph.add_run(text_after)
                                run.bold = True
                                run.font.size = Pt(11)
                                run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)
                                run.font.name = 'Poppins'
                        elif key == "CodeApprenant":
                            # Sauvegarder le texte original
                            original_text = paragraph.text
                            # Trouver la position du placeholder
                            placeholder_start = original_text.find(placeholder)
                            # Obtenir le texte avant le placeholder (qui inclut "Identifiant : ")
                            text_before = original_text[:placeholder_start]
                            # Effacer le texte du paragraphe
                            paragraph.text = ""
                            # Ajouter "Identifiant : " en blanc (invisible)
                            run = paragraph.add_run("Identifiant : ")
                            run.font.color.rgb = RGBColor(255, 255, 255)  # Blanc
                            # Ajouter la valeur en blanc (invisible)
                            run = paragraph.add_run(str(value))
                            run.font.color.rgb = RGBColor(255, 255, 255)  # Blanc
                        elif key.endswith("_Title"):
                            # Supprimer le placeholder
                            paragraph.text = paragraph.text.replace(placeholder, "")
                            # Ajouter le texte en gras
                            run = paragraph.add_run(str(value))
                            run.bold = True
                        elif key.startswith("moyUE") or key.startswith("ECTSUE") or key.startswith("etatUE"):
                            # Supprimer le placeholder
                            paragraph.text = paragraph.text.replace(placeholder, "")
                            # Ajouter le texte en gras et centré
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.add_run(str(value))
                            run.bold = True
                        elif key in ["moyenne", "moyenneECTS", "totaletat"]:
                            # Supprimer le placeholder
                            paragraph.text = paragraph.text.replace(placeholder, "")
                            # Ajouter le texte en gras et en blanc
                            run = paragraph.add_run(str(value))
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)  # Blanc
                        elif key.startswith("etat") and not key.startswith("etatUE"):
                            # Supprimer le placeholder
                            paragraph.text = paragraph.text.replace(placeholder, "")
                            run = paragraph.add_run(str(value))
                            if str(value) == "R":
                                run.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0x69, 0x59)  # #FF6959
                        elif (key.startswith("note") or key.startswith("ECTS")):
                            # Supprimer le placeholder
                            cell.text = cell.text.replace(placeholder, "")
                            # Pour les notes et ECTS, centrer le texte
                            paragraph = cell.paragraphs[0]
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrer le texte
                            run = paragraph.add_run(str(value))
                            run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)  # Couleur #0a5d81
                            run.font.name = 'Poppins'
                        elif key.startswith("Absences justifiees") or key.startswith("Absences injustifiees") or key.startswith("Retards"):
                            cell.text = cell.text.replace(placeholder, "")
                            run = cell.paragraphs[0].add_run(str(value))
                            run.font.name = 'Poppins'
                            run.font.size = Pt(8)
                            run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)
                        elif key.startswith("justifiee") or key.startswith("injustifiee") or key.startswith("retard"):
                            cell.text = cell.text.replace(placeholder, "")
                            paragraph = cell.paragraphs[0]
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.add_run(str(value))
                            run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)
                            run.font.name = 'Poppins'
                            run.font.size = Pt(8)
                        elif key == "datedujour":
                            # Sauvegarder le texte original
                            original_text = paragraph.text
                            # Trouver la position du placeholder
                            placeholder_start = original_text.find(placeholder)
                            # Obtenir le texte avant et après le placeholder
                            text_before = original_text[:placeholder_start]
                            text_after = original_text[placeholder_start + len(placeholder):]
                            
                            # Effacer le texte du paragraphe
                            paragraph.text = ""
                            # Aligner le paragraphe à droite
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            
                            # Ajouter le texte avant avec le style
                            if text_before:
                                run = paragraph.add_run(text_before)
                                run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)
                                run.font.name = 'Poppins'
                                run.font.size = Pt(8)
                            
                            # Ajouter la date avec le style
                            run = paragraph.add_run(str(value))
                            run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)
                            run.font.name = 'Poppins'
                            run.font.size = Pt(8)
                            
                            # Ajouter le texte après avec le style
                            if text_after:
                                run = paragraph.add_run(text_after)
                                run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)
                                run.font.name = 'Poppins'
                                run.font.size = Pt(8)
                        # Cas spécial pour les APPRECIATIONS
                        elif key == "APPRECIATIONS":
                            # Si la valeur est vide ou contient uniquement des espaces
                            if not value or value.strip() == "":
                                # Effacer le texte du paragraphe
                                paragraph.text = ""
                                # Ajouter le placeholder en blanc
                                run = paragraph.add_run("{{APPRECIATIONS}}")
                                run.font.color.rgb = RGBColor(255, 255, 255)  # Blanc
                            else:
                                # Sinon, remplacer normalement
                                paragraph.text = paragraph.text.replace(placeholder, str(value))
                        else:
                            # Pour les autres variables, remplacement normal
                            paragraph.text = paragraph.text.replace(placeholder, str(value))

            
            
            # Remplacer dans les tableaux
            for table in doc.tables:
                for table_row in table.rows:
                    for cell in table_row.cells:
                        for key, value in student_data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in cell.text:
                                if key == "etendugroupe":
                                    # Sauvegarder le texte original
                                    original_text = cell.text
                                    # Trouver la position du placeholder
                                    placeholder_start = original_text.find(placeholder)
                                    # Obtenir le texte après le placeholder
                                    text_after = original_text[placeholder_start + len(placeholder):]
                                    # Effacer le texte de la cellule
                                    cell.text = ""
                                    paragraph = cell.paragraphs[0]
                                    # Ajouter la valeur avec le style voulu
                                    run = paragraph.add_run(str(value))
                                    run.bold = True
                                    run.font.size = Pt(11)
                                    run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)
                                    run.font.name = 'Poppins'
                                    # Ajouter le texte qui suivait le placeholder
                                    if text_after:
                                        run = paragraph.add_run(text_after)
                                        run.bold = True
                                        run.font.size = Pt(11)
                                        run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)
                                        run.font.name = 'Poppins'
                                # Dans la section des tableaux
                                elif key == "CodeApprenant":
                                    # Sauvegarder le texte original
                                    original_text = cell.text
                                    # Trouver la position du placeholder
                                    placeholder_start = original_text.find(placeholder)
                                    # Obtenir le texte avant le placeholder (qui inclut "Identifiant : ")
                                    text_before = original_text[:placeholder_start]
                                    # Effacer le texte de la cellule
                                    cell.text = ""
                                    paragraph = cell.paragraphs[0]
                                    # Ajouter "Identifiant : " en blanc (invisible)
                                    run = paragraph.add_run("Identifiant : ")
                                    run.font.color.rgb = RGBColor(255, 255, 255)  # Blanc
                                    # Ajouter la valeur en blanc (invisible)
                                    run = paragraph.add_run(str(value))
                                    run.font.color.rgb = RGBColor(255, 255, 255)  # Blanc
                                elif key.endswith("_Title"):
                                    # Supprimer le placeholder
                                    cell.text = cell.text.replace(placeholder, "")
                                    # Ajouter le texte en gras dans le premier paragraphe de la cellule
                                    run = cell.paragraphs[0].add_run(str(value))
                                    run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)  # Couleur #0a5d81
                                    run.font.name = 'Poppins'
                                    run.font.size = Pt(8)
                                    run.bold = True
                                elif key.startswith("matiere"):
                                    # Supprimer le placeholder
                                    cell.text = cell.text.replace(placeholder, "")
                                    # Ajouter le texte en gras dans le premier paragraphe de la cellule
                                    run = cell.paragraphs[0].add_run(str(value))
                                    run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)  # Couleur #0a5d81
                                    run.font.name = 'Poppins'
                                    run.font.size = Pt(8)
                                elif key.startswith("moyUE") or key.startswith("ECTSUE") or key.startswith("etatUE"):
                                    # Supprimer le placeholder
                                    cell.text = cell.text.replace(placeholder, "")
                                    # Ajouter le texte en gras et centré
                                    paragraph = cell.paragraphs[0]
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = paragraph.add_run(str(value))
                                    run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)  # Couleur #0a5d81
                                    run.font.name = 'Poppins'
                                    run.font.size = Pt(8)
                                    run.bold = True
                                elif key in ["moyenne", "moyenneECTS", "totaletat"]:
                                    # Supprimer le placeholder
                                    cell.text = cell.text.replace(placeholder, "")
                                    # Ajouter le texte en gras et en blanc dans le premier paragraphe de la cellule
                                    paragraph = cell.paragraphs[0]
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrer le texte
                                    run = paragraph.add_run(str(value))
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)  # Blanc
                                elif key.startswith("etat") and not key.startswith("etatUE"):
                                    # Supprimer le placeholder
                                    cell.text = cell.text.replace(placeholder, "")
                                    # Pour les états, centrer le texte et colorer en rouge si "R"
                                    paragraph = cell.paragraphs[0]
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = paragraph.add_run(str(value))
                                    run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)  # Couleur #0a5d81
                                    run.font.name = 'Poppins'
                                    run.font.size = Pt(8)
                                    
                                    if str(value) == "R":
                                        run.bold = True
                                        run.font.color.rgb = RGBColor(0xFF, 0x69, 0x59)  # #FF6959
                                        run.font.name = 'Poppins'
                                        run.font.size = Pt(8)
                                        
                                elif (key.startswith("note") or key.startswith("ECTS")):
                                    # Supprimer le placeholder
                                    cell.text = cell.text.replace(placeholder, "")
                                    # Pour les notes et ECTS, centrer le texte
                                    paragraph = cell.paragraphs[0]
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrer le texte
                                    run = paragraph.add_run(str(value))
                                    run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)  # Couleur #0a5d81
                                    run.font.name = 'Poppins'
                                    run.font.size = Pt(8)
                                elif key.startswith("Absences justifiees") or key.startswith("Absences injustifiees") or key.startswith("Retards"):
                                    cell.text = cell.text.replace(placeholder, "")
                                    run = cell.paragraphs[0].add_run(str(value))
                                    run.font.name = 'Poppins'
                                    run.font.size = Pt(8)
                                    run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)
                                elif key.startswith("justifiee") or key.startswith("injustifiee") or key.startswith("retard"):
                                    cell.text = cell.text.replace(placeholder, "")
                                    paragraph = cell.paragraphs[0]
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = paragraph.add_run(str(value))
                                    run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)
                                    run.font.name = 'Poppins'
                                    run.font.size = Pt(8)
                                    # Log pour confirmer le remplacement
                                    logging.info(f"Texte final après remplacement : {paragraph.text}")
                                elif key == "datedujour":
                                    # Sauvegarder le texte original
                                    original_text = cell.text
                                    # Trouver la position du placeholder
                                    placeholder_start = original_text.find(placeholder)
                                    # Obtenir le texte avant et après le placeholder
                                    text_before = original_text[:placeholder_start]
                                    text_after = original_text[placeholder_start + len(placeholder):]
                                    
                                    # Effacer le texte de la cellule
                                    cell.text = ""
                                    paragraph = cell.paragraphs[0]
                                    
                                    # Ajouter le texte avant avec le style
                                    if text_before:
                                        run = paragraph.add_run(text_before)
                                        run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)
                                        run.font.name = 'Poppins'
                                        run.font.size = Pt(8)
                                    
                                    # Ajouter la date avec le style
                                    run = paragraph.add_run(str(value))
                                    run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)
                                    run.font.name = 'Poppins'
                                    run.font.size = Pt(8)
                                    
                                    # Ajouter le texte après avec le style
                                    if text_after:
                                        run = paragraph.add_run(text_after)
                                        run.font.color.rgb = RGBColor(0x0A, 0x5D, 0x81)
                                        run.font.name = 'Poppins'
                                        run.font.size = Pt(8)
                                elif key == "APPRECIATIONS":
                                    # Si la valeur est vide ou contient uniquement des espaces
                                    if not value or value.strip() == "":
                                        # Effacer le texte de la cellule
                                        cell.text = ""
                                        paragraph = cell.paragraphs[0]
                                        # Ajouter le placeholder en blanc
                                        run = paragraph.add_run("{{APPRECIATIONS}}")
                                        run.font.color.rgb = RGBColor(255, 255, 255)  # Blanc
                                    else:
                                        # Sinon, remplacer normalement
                                        cell.text = cell.text.replace(placeholder, str(value))
                                else:
                                    # Pour les autres variables, remplacement normal sans centrage
                                    cell.text = cell.text.replace(placeholder, str(value))

                        # Ajouter la signature si le code apprenant correspond
            
                        # Ajouter la signature si le code apprenant correspond

            # Sauvegarder le bulletin
            safe_nom = "".join(c for c in student_data["nomApprenant"] if c.isalnum() or c in (' ', '-', '_')).strip()
            bulletin_path = os.path.join(bulletins_dir, f"bulletin_{safe_nom}.docx")
            doc.save(bulletin_path)
            logging.info(f"Bulletin créé pour {student_data['nomApprenant']}")

        await db.disconnect()
        return {
            "message": "Bulletins générés avec succès",
            "bulletins_directory": bulletins_dir
        }

    except Exception as e:
        logging.error(f"Erreur lors de la génération des bulletins : {str(e)}")
        raise HTTPException(status_code=400, detail=str(e))
