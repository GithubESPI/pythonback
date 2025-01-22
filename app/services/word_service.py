import base64
import os
import logging
import zipfile
from docx import Document
from zipfile import ZipFile
import openpyxl
from io import BytesIO
from app.services.ects_service import get_ects_for_template
from prisma import Prisma
from app.services.prisma_service import get_excel_from_prisma, get_template_from_prisma

async def save_word_template(template_name: str, output_dir: str) -> str:
    """
    Sauvegarde le template Word depuis Prisma vers le dossier temporaire.
    """
    try:
        db = Prisma()
        await db.connect()
        
        word_template = await db.generatedfile.find_first(
            where={
                "filename": template_name,
                "isTemplate": True
            }
        )
        
        if not word_template:
            raise ValueError(f"Template Word '{template_name}' non trouvé")
            
        # Créer le dossier de sortie s'il n'existe pas
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            
        # Sauvegarder le template Word
        template_path = os.path.join(output_dir, template_name)
        with open(template_path, "wb") as f:
            f.write(word_template.fileData)
            
        await db.disconnect()
        return template_path
        
    except Exception as e:
        logging.error(f"Erreur lors de la sauvegarde du template Word : {str(e)}")
        raise

async def generate_bulletins_from_excel(excel_id: int, output_dir: str):
    try:
        logging.info(f"Début de la génération des bulletins pour l'Excel ID: {excel_id}")
        
        db = Prisma()
        await db.connect()
        
        # Récupérer l'Excel généré
        generated_excel = await db.generatedexcel.find_unique(
            where={"id": excel_id},
            include={"template": True}
        )
        
        if not generated_excel:
            raise ValueError(f"Excel généré non trouvé avec l'ID : {excel_id}")

        # Charger l'Excel
        excel_bytes = base64.b64decode(str(generated_excel.data))
        excel_wb = openpyxl.load_workbook(BytesIO(excel_bytes))
        excel_ws = excel_wb.active

        # Déterminer le template Word à utiliser en fonction du groupe
        group_name = str(excel_ws["B2"].value or "").strip()
        template_name = "modeleBGALT3.docx"
        
        if "BG-ALT-2" in group_name or "BG-ALT-S2" in group_name:
            template_name = "modeleBGALT2.docx"
            logging.info(f"Utilisation du template {template_name} pour le groupe {group_name}")
        else:
            logging.info(f"Utilisation du template par défaut {template_name} pour le groupe {group_name}")

        # Récupérer le template Word
        word_template = await db.generatedfile.find_first(
            where={
                "filename": template_name,
                "isTemplate": True
            }
        )
        
        if not word_template:
            raise ValueError(f"Template Word {template_name} non trouvé")

        # Créer le dossier pour les bulletins
        bulletins_dir = os.path.join(output_dir, "bulletins")
        if not os.path.exists(bulletins_dir):
            os.makedirs(bulletins_dir)

        # Sauvegarder temporairement le template Word
        temp_word_path = os.path.join(output_dir, template_name)
        word_bytes = base64.b64decode(str(word_template.fileData))
        with open(temp_word_path, 'wb') as f:
            f.write(word_bytes)

        # Configuration des colonnes selon le template
        if template_name == "modeleBGALT2.docx":
            column_config = {
                "nom_prenom": "B",
                "date_naissance": "V",
                "site": "W",
                "code_groupe": "Y",
                "nom_groupe": "Z",
                "abs_justifiees": "AA",
                "abs_injustifiees": "AB",
                "retards": "AC",
                "appreciation": "AD"
            }
        else:  # modeleBGALT3.docx
            column_config = {
                "nom_prenom": "B",
                "date_naissance": "T",
                "site": "U",
                "code_groupe": "W",
                "nom_groupe": "X",
                "abs_justifiees": "Y",
                "abs_injustifiees": "Z",
                "retards": "AA",
                "appreciation": "AB"
            }

        # Pour chaque étudiant, créer un bulletin personnalisé
        for row in range(3, excel_ws.max_row + 1):
            if not excel_ws[f"B{row}"].value:
                continue

            try:
                doc = Document(temp_word_path)
                
                # Préparer les données de l'étudiant selon la configuration
                student_data = {
                    "NOM_PRENOM": excel_ws[f"{column_config['nom_prenom']}{row}"].value or "",
                    "DATE_NAISSANCE": excel_ws[f"{column_config['date_naissance']}{row}"].value or "",
                    "SITE": excel_ws[f"{column_config['site']}{row}"].value or "",
                    "CODE_GROUPE": excel_ws[f"{column_config['code_groupe']}{row}"].value or "",
                    "NOM_GROUPE": excel_ws[f"{column_config['nom_groupe']}{row}"].value or "",
                    "ABS_JUSTIFIEES": str(excel_ws[f"{column_config['abs_justifiees']}{row}"].value or ""),
                    "ABS_INJUSTIFIEES": str(excel_ws[f"{column_config['abs_injustifiees']}{row}"].value or ""),
                    "RETARDS": str(excel_ws[f"{column_config['retards']}{row}"].value or ""),
                    "APPRECIATION": excel_ws[f"{column_config['appreciation']}{row}"].value or ""
                }

                # Remplacer les variables dans le document
                for paragraph in doc.paragraphs:
                    for key, value in student_data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, str(value))

                # Remplacer dans les tableaux
                for table in doc.tables:
                    for table_row in table.rows:
                        for cell in table_row.cells:
                            for key, value in student_data.items():
                                placeholder = f"{{{{{key}}}}}"
                                if placeholder in cell.text:
                                    cell.text = cell.text.replace(placeholder, str(value))

                # Sauvegarder le bulletin
                safe_nom = "".join(c for c in student_data["NOM_PRENOM"] if c.isalnum() or c in (' ', '-', '_')).strip()
                bulletin_path = os.path.join(bulletins_dir, f"bulletin_{safe_nom}.docx")
                doc.save(bulletin_path)
                logging.info(f"Bulletin créé pour {student_data['NOM_PRENOM']}")

            except Exception as e:
                logging.error(f"Erreur lors du traitement de l'étudiant à la ligne {row}: {str(e)}")
                continue

        # Nettoyer le fichier temporaire
        if os.path.exists(temp_word_path):
            os.remove(temp_word_path)

        await db.disconnect()
        return bulletins_dir

    except Exception as e:
        logging.error(f"Erreur lors de la génération des bulletins : {str(e)}")
        raise
